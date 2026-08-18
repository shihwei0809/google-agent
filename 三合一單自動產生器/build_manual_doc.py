import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import win32com.client

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_manual_docx(docx_path):
    doc = docx.Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 標題
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("📘 三合一單與運輸通知表自動產生器\n操作手冊 (User Manual)")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
    p_title.paragraph_format.space_after = Pt(15)

    # 副標題
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_sub = p_sub.add_run("適用對象：現場、排程與物流作業人員 | 版本：v2.0")
    r_sub.font.size = Pt(9.5)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p_sub.paragraph_format.space_after = Pt(20)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Microsoft JhengHei'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Microsoft JhengHei'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            r_b.font.bold = True
            r_b.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
        p.add_run(text)
        return p

    def add_tip(text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "E8F5E9")
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("💡 " + text)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 1. 系統簡介
    add_h1("一、 系統簡介")
    p1 = doc.add_paragraph("本工具專為物流、排程與現場作業人員設計，提供簡單易用的桌面圖形介面 (GUI)。能夠根據「批號」與「送達地點」，一鍵批次完成以下兩大報表的自動化產生：")
    p1.paragraph_format.space_after = Pt(6)
    add_bullet("自動寫入槽號、批號、送達地點，並重新產生獨立條碼 (QR Code) 圖片。", "1. 獨立三合一單 Excel：")
    add_bullet("彙整所有項目為精美排程卡片，支援預計到廠時間、修正到廠時間刪除線特效與併排修正通知卡片。", "2. 運輸通知表 Excel：")

    # 2. 一鍵啟動與系統檢查
    add_h1("二、 一鍵啟動與系統檢查")
    add_h2("1. 快速啟動方式")
    p2 = doc.add_paragraph("在專案資料夾中，直接雙擊執行：")
    p2.paragraph_format.space_after = Pt(4)
    p_cmd = doc.add_paragraph()
    p_cmd.paragraph_format.left_indent = Inches(0.3)
    r_cmd = p_cmd.add_run("👉 啟動_三合一單產生器.bat")
    r_cmd.font.bold = True
    r_cmd.font.size = Pt(12)
    r_cmd.font.color.rgb = RGBColor(0x21, 0x96, 0xF3)
    p_cmd.paragraph_format.space_after = Pt(6)
    add_tip("工具執行時會自動偵測 Python 環境，若缺少 openpyxl、qrcode 或 pillow 套件，將會自動為您完成安裝並開啟視窗，無需手動設定。")

    add_h2("2. 系統狀態檢查")
    doc.add_paragraph("軟體開啟後，頂部「系統狀態」區會自動核對以下兩個必備檔案：")
    add_bullet("台積電槽車barcode三合一單-範本.xlsx", "• 範本檔案：")
    add_bullet("地點代號對照表.xlsx（顯示已載入之地點筆數）", "• 對照表檔案：")

    # 3. 介面功能與操作說明
    add_h1("三、 介面功能與操作說明")
    add_h2("1. 頂部「一鍵批次設定」列")
    doc.add_paragraph("為方便快速填寫相同資訊，頂部提供批次工具列：")
    add_bullet("一鍵勾選或取消勾選下方所有資料列。", "☑️ 全選 / 全不選：")
    add_bullet("可直接輸入日期，或點擊 📅 小圖示開啟日曆視窗點選日期，點擊 [套用至全列] 將日期寫入所有有效資料列。", "📅 批次日期：")
    add_bullet("輸入預計到達時間 (如 10:00)，點擊 [套用至全列] 套用。", "⏰ 批次預計時間：")
    add_bullet("輸入修改後的到達時間 (如 08:00)，點擊 [套用至全列] 套用。", "🔴 批次修正時間：")
    add_bullet("可勾選是否產生「三合一單」與/或「運輸通知表」 (預設均為勾選)。", "📊 報表種類開關：")

    add_h2("2. 表格欄位詳細說明")
    
    table_spec = doc.add_table(rows=1, cols=3)
    table_spec.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_spec.autofit = False
    
    hdr_cells = table_spec.rows[0].cells
    hdr_titles = ["欄位名稱", "功能說明", "備註 / 特色"]
    hdr_widths = [Inches(1.5), Inches(3.2), Inches(2.3)]
    for i, title in enumerate(hdr_titles):
        cell = hdr_cells[i]
        cell.width = hdr_widths[i]
        set_cell_background(cell, "002060")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rows_data = [
        ("產生", "勾選是否將該筆列入報表產出", "預設均打勾 ☑️"),
        ("項次", "資料順序編號 (1, 2, 3...)", "自動編號"),
        ("批號", "輸入剛好 10 碼的批號", "輸入後自動帶出 5~8 碼槽號"),
        ("槽號", "依批號規則自動帶出之槽號", "藍字顯示 (唯讀)"),
        ("地點", "輸入送達地點 (如 15P5 或 18P3B)", "輸入後自動轉換長代號"),
        ("長代號", "對照表轉換後之完整代號", "紫字顯示，未知顯示 ❌"),
        ("出貨日期 📅", "出貨日期，點擊 📅 圖示開啟日曆", "預設空白，附帶月曆選擇器"),
        ("預計到廠時間", "輸入預計到達時間 (如 10:00)", "預設空白"),
        ("修正到廠時間", "輸入修改後的到達時間 (如 08:00)", "紅字顯示，自動觸發刪除線")
    ]

    for row_idx, data in enumerate(rows_data):
        row_cells = table_spec.add_row().cells
        bg_color = "F9F9F9" if row_idx % 2 == 1 else "FFFFFF"
        for i, text in enumerate(data):
            cell = row_cells[i]
            cell.width = hdr_widths[i]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(text)
                r.font.bold = True
            else:
                p.add_run(text)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 4. 智慧貼上與 Excel 匯入
    add_h1("四、 智慧貼上與 Excel 匯入技巧")
    add_h2("1. 智慧 Ctrl+V 快速複製貼上")
    doc.add_paragraph("您可以直接從 Excel 複製多筆排程資料，並點擊軟體第一列的「批號」欄位按下 Ctrl+V：")
    add_bullet("系統會自動分割 Tab 欄位，將資料依序填入批號與地點，並自動計算槽號與長代號。", "• 複製 2 欄資料 (批號 + 地點)：")
    add_bullet("系統可一次自動帶入批號、地點、日期、預計時間與修正時間。", "• 複製 3~5 欄資料：")

    add_h2("2. 📥 從 Excel 檔案直接匯入")
    doc.add_paragraph("點擊軟體右上方藍色按鈕 `📥 從 Excel 匯入`：")
    add_bullet("選擇任意包含排程的 Excel 檔案，系統會自動尋找包含「批號」、「地點/送達地點」、「日期」、「時間」之欄位並自動填入表格中。")

    # 5. 報表產出與檔名規格
    add_h1("五、 報表產出與檔名規格")
    doc.add_paragraph("點擊下方 `🚀 開始批次產生 Excel 報表` 後，系統將在與程式同目錄下建立當天日期的歸檔資料夾：")
    
    add_h2("1. 當天日期歸檔資料夾")
    add_bullet("三合一單輸出_YYYYMMDD (例如：三合一單輸出_20260818)", "📁 資料夾名稱：")
    add_tip("一天中多次產生的所有檔案皆會集中存放在當天的資料夾內，免去產生大量帶時間戳記的重複資料夾，抓檔輕鬆便利！")

    add_h2("2. 三合一單獨立 Excel 檔名規格")
    add_bullet("[出貨日期]. [地點]台積電槽車barcode三合一單.xlsx", "📄 檔名命名公式：")
    add_bullet("2026.8.18. 18P3B台積電槽車barcode三合一單.xlsx", "• 範例 1：出貨日期 2026/8/18，地點 18P3B ➔ ")
    add_bullet("2026.8.19. 15P5台積電槽車barcode三合一單.xlsx", "• 範例 2：出貨日期 2026/8/19，地點 15P5 ➔ ")
    doc.add_paragraph("檔名完全依據出貨日期與地點自動組合，方便現場人員迅速核對與存取！")

    add_h2("3. 運輸通知表 Excel (運輸通知表.xlsx)")
    add_bullet("在 A~F 欄生成標準「出貨排程通知」卡片，預計到廠時間以藍字加粗顯示。", "• 無修正時間時：")
    add_bullet("在 A~F 欄顯示原始卡片，同時在同分頁右側 (H~M 欄) 自動併排生成「出貨排程修正通知」卡片！第一個預計時間自動標上刪除線 (~10:00~)，修正時間填寫於下方並以紅字加粗強調。", "• 有填寫修正時間時：")

    # 6. FAQ
    add_h1("六、 常問問題與故障排除 (FAQ)")
    add_bullet("請確認批號是否剛好為 10 碼。若長度不符，請檢查是否有贅字或少打。", "Q1：跳出「批號長度錯誤」警示？ ➔ ")
    add_bullet("表示輸入之地點代號尚未登錄在『地點代號對照表.xlsx』中。請開啟對照表檔新增對應之地點與長代號後儲存，再重新執行即可。", "Q2：跳出「地點代號對照表中找不到」？ ➔ ")
    add_bullet("點擊日期欄位旁邊的 📅 圖示即可開啟月曆視窗。點選欲設定之日期或點擊 [今天] / [明天] 鈕即可完成輸入。", "Q3：如何開啟與使用日期選擇器？ ➔ ")

    doc.save(docx_path)
    print(f"Word file generated: {docx_path}")

def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        # 17 represents wdFormatPDF
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        word.Quit()
        print(f"PDF file generated: {pdf_path}")
    except Exception as e:
        print(f"Failed to convert to PDF via Word COM: {e}")

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.abspath(__file__))
    docx_file = os.path.join(base_dir, "三合一單與運輸通知表自動產生器_操作手冊.docx")
    pdf_file = os.path.join(base_dir, "三合一單與運輸通知表自動產生器_操作手冊.pdf")
    
    create_manual_docx(docx_file)
    convert_docx_to_pdf(docx_file, pdf_file)
