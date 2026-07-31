"""
消防演習影音分鏡腳本與提示詞產生器
後端 API (FastAPI)
"""
import json
import logging
import os
from pathlib import Path
from typing import List, Optional
import requests
import fitz  # PyMuPDF
import docx  # python-docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from urllib.parse import quote
from fastapi import FastAPI, Form, HTTPException, Request, UploadFile, File
from fastapi.responses import HTMLResponse, StreamingResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

# ─── Log & Configuration ──────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

app = FastAPI(title="消防演習影音分鏡腳本與提示詞產生器")

# Templates and Static Files
BASE_DIR = Path(__file__).resolve().parent
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))

# Define folder structure
(BASE_DIR / "templates").mkdir(exist_ok=True)

# ─── Pydantic Models for API ──────────────────────────────────────────────────
class Scene(BaseModel):
    scene_no: int
    phase: str
    narrator_script: str
    action_description: str
    visual_concept: str
    motion_concept: str
    image_prompt: Optional[str] = None
    video_prompt: Optional[str] = None

class StoryboardRequest(BaseModel):
    drill_text: str
    api_key: str
    model: str = "gemini-3.5-flash"

class PromptRequest(BaseModel):
    scenes: List[Scene]
    api_key: str
    model: str = "gemini-3.5-flash"

# ─── Gemini & Grok JSON Schema Definitions ─────────────────────────────────────
STORYBOARD_SCHEMA = {
    "type": "OBJECT",
    "properties": {
        "scenes": {
            "type": "ARRAY",
            "items": {
                "type": "OBJECT",
                "properties": {
                    "scene_no": {"type": "INTEGER"},
                    "phase": {"type": "STRING"},
                    "narrator_script": {"type": "STRING"},
                    "action_description": {"type": "STRING"},
                    "visual_concept": {"type": "STRING"},
                    "motion_concept": {"type": "STRING"}
                },
                "required": ["scene_no", "phase", "narrator_script", "action_description", "visual_concept", "motion_concept"]
            }
        }
    },
    "required": ["scenes"]
}

# ─── Gemini & Grok JSON Schema Definitions ─────────────────────────────────────
PROMPTS_SCHEMA = {
    "type": "OBJECT",
    "properties": {
        "prompts": {
            "type": "ARRAY",
            "items": {
                "type": "OBJECT",
                "properties": {
                    "scene_no": {"type": "INTEGER"},
                    "image_prompt": {"type": "STRING"},
                    "video_prompt": {"type": "STRING"}
                },
                "required": ["scene_no", "image_prompt", "video_prompt"]
            }
        }
    },
    "required": ["prompts"]
}

# ─── Gemini & Grok API Call Helpers ───────────────────────────────────────────
def call_gemini_json(api_key: str, model: str, system_instruction: str, user_prompt: str, response_schema: Optional[dict] = None) -> dict:
    """Call Google Gemini API with JSON response format enabled and optional schema validation."""
    try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
        headers = {"Content-Type": "application/json"}
        
        payload = {
            "contents": [
                {
                    "role": "user",
                    "parts": [{"text": user_prompt}]
                }
            ],
            "systemInstruction": {
                "parts": [{"text": system_instruction}]
            },
            "generationConfig": {
                "responseMimeType": "application/json"
            }
        }
        
        if response_schema:
            payload["generationConfig"]["responseSchema"] = response_schema
        
        response = requests.post(url, json=payload, headers=headers, timeout=90)
        response.raise_for_status()
        data = response.json()
        
        # Extract response text
        text_content = data["candidates"][0]["content"]["parts"][0]["text"]
        logger.info("Successfully received response from Gemini API.")
        return json.loads(text_content.strip())
    except Exception as e:
        logger.exception("Gemini API request failed")
        raise HTTPException(status_code=500, detail=f"呼叫 Gemini API 失敗：{str(e)}")

def call_grok_json(api_key: str, model: str, system_instruction: str, user_prompt: str) -> dict:
    """Call xAI Grok API with JSON response format enabled."""
    try:
        url = "https://api.xai.ai/v1/chat/completions"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        
        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": system_instruction},
                {"role": "user", "content": user_prompt}
            ],
            "response_format": {"type": "json_object"}
        }
        
        response = requests.post(url, json=payload, headers=headers, timeout=90)
        response.raise_for_status()
        data = response.json()
        
        text_content = data["choices"][0]["message"]["content"]
        logger.info("Successfully received response from Grok API.")
        return json.loads(text_content.strip())
    except Exception as e:
        logger.exception("Grok API request failed")
        raise HTTPException(status_code=500, detail=f"呼叫 Grok API 失敗：{str(e)}")

# ─── Routes ───────────────────────────────────────────────────────────────────
@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse(request=request, name="index.html")

@app.post("/api/import-file")
async def import_file(file: UploadFile = File(...)):
    filename = file.filename.lower()
    content = await file.read()
    
    try:
        if filename.endswith(".txt"):
            text = content.decode("utf-8")
        elif filename.endswith(".pdf"):
            doc = fitz.open(stream=content, filetype="pdf")
            text = "\n".join([page.get_text() for page in doc]).strip()
        elif filename.endswith(".docx"):
            doc = docx.Document(BytesIO(content))
            paragraphs = [para.text for para in doc.paragraphs]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    paragraphs.append(" | ".join(row_text))
            text = "\n".join(paragraphs).strip()
        else:
            raise HTTPException(status_code=400, detail="不支援的檔案格式。僅支援 .txt、.pdf、.docx")
            
        if not text.strip():
            raise HTTPException(status_code=400, detail="讀取成功，但檔案中未發現任何可解析的文字內容。")
            
        return {"text": text}
    except Exception as e:
        logger.exception("File import failed")
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=f"檔案解析失敗：{str(e)}")

@app.post("/api/generate-storyboard")
async def generate_storyboard(req: StoryboardRequest):
    if not req.drill_text.strip():
        raise HTTPException(status_code=400, detail="演習內容不可為空。")
    if not req.api_key.strip():
        raise HTTPException(status_code=400, detail="請填寫 API Key。")
        
    system_instruction = (
        "你是一個專業的化學工廠安全演習導演與影音分鏡設計師。你的任務是將使用者提供的消防演練或災害應變計畫內容，"
        "拆解並撰寫成一組連續且極富視覺張力的「影音分鏡腳本」。\n\n"
        "請遵循以下規則：\n"
        "1. 依故事進度順序拆解為更細緻的 10 到 20 個分鏡卡片（視計畫內容豐富度可彈性增加分鏡數量）。\n"
        "2. **單一分鏡焦點原則**：每個分鏡卡片必須「僅專注於單一的實體動作、通報項目或演練事件」。絕對不要在同一個分鏡中塞入兩個或多個不同的演練項目（例如：不要將『林東和關閉馬達閥門』與『陳世偉聯絡警衛廣播』合併；不要將『葉豪恩關閉 MCC 電力』與『陳又詳啟動泡沫槽』合併；不要將『一班滅火』與『二班降溫』合併）。每個獨立的行動步驟都應該擁有自己專屬的分鏡卡片。這樣做能確保每個分鏡在生成 AI 圖片和影片時，畫面焦點單純且細節豐富，避免畫面空洞、混亂或語意不清。\n"
        "3. 每個分鏡應包含：\n"
        "   - `scene_no`: 數字，分鏡序號 (從 1 開始遞增)。\n"
        "   - `phase`: 字串，目前的應變演習階段（例如：事故起因、緊急通報、初期滅火、成立指揮所、專業防護、救援警戒、善後處理）。\n"
        "   - `narrator_script`: 字串，沉穩、清晰、專業的廣播或旁白中文配音稿。字數約 80-150 字，適合語音朗讀。\n"
        "   - `action_description`: 字串，簡短描述現場人員應進行的單一實體操作與動作（如：拉起隔離帶、啟動泡沫閥、搬運吸液棉等）。\n"
        "   - `visual_concept`: 字串，畫面視覺構想。請以繁體中文詳細描述該場景的畫面細節，包含地點、人事物、消防裝備、火焰、煙霧及光影氛圍，以利後續生成圖片。請確保只描繪一個焦點畫面的視覺細節。\n"
        "   - `motion_concept`: 字串，鏡頭運鏡與動態構想。請以繁體中文描述鏡頭的動態（如：鏡頭緩慢向右橫移、從低角度往上仰拍、水柱噴射的慢動作、煙霧升騰軌跡等），以利生成影片。\n\n"
        "請務必輸出 JSON 格式，且完全依照規定的欄位名稱。不要有任何 Markdown 包裹文字。"
    )
    
    user_prompt = (
        f"請將以下消防演習計畫的內容轉換為分鏡腳本：\n\n{req.drill_text}\n\n"
        "請嚴格回傳以下格式的 JSON 物件：\n"
        "{\n"
        '  "scenes": [\n'
        "    {\n"
        '      "scene_no": 1,\n'
        '      "phase": "階段名稱",\n'
        '      "narrator_script": "配音稿內容",\n'
        '      "action_description": "動作說明",\n'
        '      "visual_concept": "畫面視覺細節構想",\n'
        '      "motion_concept": "鏡頭動態運鏡構想"\n'
        "    }\n"
        "  ]\n"
        "}"
    )
    
    if req.model.startswith("grok-"):
        res = call_grok_json(req.api_key, req.model, system_instruction, user_prompt)
    else:
        res = call_gemini_json(req.api_key, req.model, system_instruction, user_prompt, STORYBOARD_SCHEMA)
        
    if "scenes" not in res:
        raise HTTPException(status_code=500, detail="API 回傳的格式不正確，缺少 'scenes' 欄位。")
    return res

@app.post("/api/generate-prompts")
async def generate_prompts(req: PromptRequest):
    if not req.scenes:
        raise HTTPException(status_code=400, detail="分鏡腳本不可為空。")
    if not req.api_key.strip():
        raise HTTPException(status_code=400, detail="請填寫 API Key。")
        
    system_instruction = (
        "你是一個頂尖的 AI 圖像與影片生成提示詞專家（Midjourney, Stable Diffusion, Runway Gen-3, Luma Dream Machine）。\n"
        "你將獲得一組已經過人員校正的分鏡腳本，你需要為每一個分鏡量身打造專門的英文提示詞：\n\n"
        "1. `image_prompt`: 專為 Midjourney V6 設計的精緻圖片提示詞。必須為英文。描述場景、主體、服裝、燈光、氛圍，"
        "並加入真實質感與相機參數（例如：photo of a chemical plant storage area, a huge silver tank labeled TK-658, a chemical tanker truck parked nearby, "
        "industrial atmosphere, heavy white smoke emitting, firemen in reflective gear, cinematic lighting, photorealistic, 8k, --ar 16:9 --style raw）。"
        "請保持簡潔有力、充滿細節且不包含引號。\n"
        "2. `video_prompt`: 專為 Runway Gen-3 或 Luma Dream Machine 設計的英文影片動態提示詞。著重於鏡頭運動、物理動態效果、煙霧、水花、人體動作等細節描述"
        "（例如：camera pans slowly from left to right, wide shot of the chemical plant storage tank, water mist spraying from the deluge system, "
        "reflective water droplets in the air, hyper-realistic, dynamic motion, 4k）。\n\n"
        "請務必輸出 JSON 格式。請嚴格維持原有的 scene_no，不要遺漏任何一個分鏡。"
    )
    
    # Structure scenes for the prompt
    scenes_text = ""
    for s in req.scenes:
        scenes_text += (
            f"--- 分鏡 {s.scene_no} ---\n"
            f"階段：{s.phase}\n"
            f"動作：{s.action_description}\n"
            f"視覺構想：{s.visual_concept}\n"
            f"運鏡與動態：{s.motion_concept}\n\n"
        )
        
    user_prompt = (
        f"請根據以下各分鏡細節，生成對應的 Midjourney 圖片提示詞 (image_prompt) 與 Runway 影片提示詞 (video_prompt)：\n\n"
        f"{scenes_text}\n"
        "請嚴格回傳以下格式的 JSON 物件：\n"
        "{\n"
        '  "prompts": [\n'
        "    {\n"
        '      "scene_no": 1,\n'
        '      "image_prompt": "Midjourney prompt...",\n'
        '      "video_prompt": "Runway/Luma prompt..."\n'
        "    }\n"
        "  ]\n"
        "}"
    )
    
    if req.model.startswith("grok-"):
        res = call_grok_json(req.api_key, req.model, system_instruction, user_prompt)
    else:
        res = call_gemini_json(req.api_key, req.model, system_instruction, user_prompt, PROMPTS_SCHEMA)
        
    if "prompts" not in res:
        raise HTTPException(status_code=500, detail="API 回傳的格式不正確，缺少 'prompts' 欄位。")
        
    # Map prompts back to the scenes
    prompt_map = {p["scene_no"]: p for p in res["prompts"]}
    
    updated_scenes = []
    for s in req.scenes:
        mapped = prompt_map.get(s.scene_no, {})
        s.image_prompt = mapped.get("image_prompt", "")
        s.video_prompt = mapped.get("video_prompt", "")
        updated_scenes.append(s)
        
    return {"scenes": updated_scenes}

@app.post("/api/export-markdown")
async def export_markdown(req: List[Scene]):
    md_content = "# 消防演習影音分鏡腳本與 AI 提示詞計畫書\n\n"
    md_content += "本計畫書由消防演習影音分鏡腳本產生器自動生成，包含分鏡配音稿、動作說明、以及用於生成圖片/影片的 AI 提示詞。\n\n"
    
    for s in req:
        md_content += f"## 🎬 分鏡 {s.scene_no:02d}：{s.phase}\n\n"
        md_content += f"- **現場動作**：{s.action_description}\n"
        md_content += f"- **廣播/配音旁白稿**：\n  > {s.narrator_script}\n\n"
        md_content += f"- **視覺細節構想**：{s.visual_concept}\n"
        md_content += f"- **鏡頭運鏡構想**：{s.motion_concept}\n\n"
        
        if s.image_prompt:
            md_content += f"🖼️ **Midjourney 圖片生成提示詞**：\n```text\n{s.image_prompt}\n```\n\n"
        if s.video_prompt:
            md_content += f"🎥 **Runway/Luma 影片生成提示詞**：\n```text\n{s.video_prompt}\n```\n\n"
        md_content += "---\n\n"
        
    filename = "fire_drill_storyboard_prompts.md"
    
    return StreamingResponse(
        iter([md_content.encode("utf-8")]),
        media_type="text/markdown",
        headers={
            "Content-Disposition": f"attachment; filename={filename}",
            "Cache-Control": "no-cache"
        }
    )

@app.post("/api/export-docx")
async def export_docx(req: List[Scene]):
    try:
        doc = docx.Document()
        
        # Set title
        title = doc.add_paragraph()
        title_run = title.add_run("消防演習影音分鏡腳本與 AI 提示詞計畫書")
        title_run.font.name = "Microsoft JhengHei"
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set description
        desc = doc.add_paragraph()
        desc_run = desc.add_run("本計畫書包含各分鏡之配音稿、動作說明、視覺畫面設計與 AI 生成提示詞（Midjourney & Runway/Luma）。")
        desc_run.font.name = "Microsoft JhengHei"
        desc_run.font.size = Pt(10.5)
        desc_run.font.italic = True
        
        doc.add_paragraph() # spacing
        
        for s in req:
            # Heading 2 for scene
            h = doc.add_paragraph()
            h_run = h.add_run(f"🎬 分鏡 {s.scene_no:02d}：{s.phase}")
            h_run.font.name = "Microsoft JhengHei"
            h_run.font.size = Pt(16)
            h_run.font.bold = True
            h_run.font.color.rgb = RGBColor(59, 130, 246)
            
            # Action description
            p_action = doc.add_paragraph()
            p_action.add_run("• 現場動作：").bold = True
            p_action.add_run(s.action_description)
            
            # Narrator Script
            p_script = doc.add_paragraph()
            p_script.paragraph_format.left_indent = Pt(18)
            p_script.add_run("• 廣播/配音旁白稿：\n").bold = True
            quote_run = p_script.add_run(f"「{s.narrator_script}」")
            quote_run.font.italic = True
            quote_run.font.color.rgb = RGBColor(107, 114, 128)
            
            # Visual concept
            p_visual = doc.add_paragraph()
            p_visual.add_run("• 畫面視覺細節構想：").bold = True
            p_visual.add_run(s.visual_concept)
            
            # Motion concept
            p_motion = doc.add_paragraph()
            p_motion.add_run("• 鏡頭運鏡構想：").bold = True
            p_motion.add_run(s.motion_concept)
            
            # Image Prompt
            if s.image_prompt:
                p_img = doc.add_paragraph()
                p_img.add_run("🖼️ Midjourney 圖片提示詞：\n").bold = True
                img_run = p_img.add_run(s.image_prompt)
                img_run.font.name = "Consolas"
                img_run.font.size = Pt(9.5)
                img_run.font.color.rgb = RGBColor(16, 185, 129)
                
            # Video Prompt
            if s.video_prompt:
                p_vid = doc.add_paragraph()
                p_vid.add_run("🎥 Runway/Luma 影片提示詞：\n").bold = True
                vid_run = p_vid.add_run(s.video_prompt)
                vid_run.font.name = "Consolas"
                vid_run.font.size = Pt(9.5)
                vid_run.font.color.rgb = RGBColor(245, 158, 11)
                
            # Separator line
            doc.add_paragraph("─" * 45)
            
        # Ensure Chinese font style applies correctly
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.name != "Consolas":
                    run.font.name = "Microsoft JhengHei"
                    
        # Save to BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        filename = "fire_drill_storyboard_prompts.docx"
        encoded_filename = quote(filename)
        
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
                "Cache-Control": "no-cache"
            }
        )
    except Exception as e:
        logger.exception("Word export failed")
        raise HTTPException(status_code=500, detail=f"匯出 Word 失敗：{str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8003, reload=False)
