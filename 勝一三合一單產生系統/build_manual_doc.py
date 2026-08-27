import os
import win32com.client
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

base_dir = os.path.dirname(os.path.abspath(__file__))
assets_dir = os.path.join(base_dir, 'manual_assets')

def set_cell_background(cell, fill_color):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(r'<w:tcMar {}>'
                      r'<w:top w:w="{}" w:type="dxa"/>'
                      r'<w:left w:w="{}" w:type="dxa"/>'
                      r'<w:bottom w:w="{}" w:type="dxa"/>'
                      r'<w:right w:w="{}" w:type="dxa"/>'
                      r'</w:tcMar>'.format(nsdecls('w'), top, left, bottom, right))
    tcPr.append(tcMar)

def create_manual_docx(docx_path):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title - Clean without emojis
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_after = Pt(10)
    r0 = p0.add_run("勝一三合一單產生系統\n圖文步驟操作手冊 (Illustrated User Manual)")
    r0.font.name = 'Microsoft JhengHei'
    r0.font.size = Pt(22)
    r0.font.bold = True
    r0.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    # Subtitle
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1.paragraph_format.space_after = Pt(15)
    r1 = p1.add_run("適用對象：勝一現場、排程與物流作業人員 | 最新版本：v3.2 (含台積電過濾、當天至+2天優先排程、多產品品名與固定重量自動對照)")
    r1.font.name = 'Microsoft JhengHei'
    r1.font.size = Pt(9.5)
    r1.font.italic = True
    r1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Section 1 Header
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(14)
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run("一、 系統簡介與核心特色")
    r2.font.name = 'Microsoft JhengHei'
    r2.font.size = Pt(15)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    # Section 1 Desc
    p3 = doc.add_paragraph()
    r3 = p3.add_run("本系統專為勝一化工物流、排程與現場作業人員設計，以台積電槽車出貨排程與「勝一訂單」為核心，提供直觀的桌面圖形介面 (GUI)，能自動完成勝一訂單解析、純台積電排程智慧過濾、當天至+2天優先預選、多產品品名辨識與固定充填重量自動對照、10~11 碼批號槽號提取、QR Code 生成、COA 檢驗截圖 OCR 辨識貼入與併排修正通知卡片產出。")
    r3.font.name = 'Microsoft JhengHei'

    # Section 1 Bullets
    bullets = [
        ("1. 勝一訂單自動解析與純台積電智慧過濾：", "自動讀取「勝一訂單」空白班表分頁，智慧過濾排除非台積電客戶，只保留台積電出貨排程；自動將「台積12廠P1」等文字正規化為「12P1」並自動匹配轉換長代號。"),
        ("2. 「當天至+2天」智慧優先推薦：", "匯入時自動鎖定「今天、明天、後天 (D ~ D+2)」的 3 天出貨區間並自動預選，一鍵快速帶出最即時的排程資料。"),
        ("3. 多產品品名動態辨識與固定重量自動對照：", "自動辨識勝一多樣產品品名（如 SEP73E5, PMAHQ, SEP73E4, CPNE4R, IPA 等），運輸通知表卡片標題自動動態顯示 Shiny [品名] Lorry，並由系統內建對照表自動填入對應產品之固定充填重量（如 4300 KG），不需人員手動輸入重量。"),
        ("4. 10~11 碼批號槽號提取與專屬檔名：", "支援 10~11 碼勝一批號規則，自動提取中間槽號（如 S96, S125, S405），三合一單檔名自動加入槽號（如 2026.8.28. 15P7_S405_台積電槽車barcode三合一單.xlsx），徹底防呆防覆蓋。"),
        ("5. COA 檢驗智慧 OCR 辨識與裁切貼入：", "支援一鍵上傳檢驗截圖，系統自動以 OCR 解析批號並精確裁切檢驗表頭與數據列，自動嵌入三合一單 F5 儲存格，免除人工剪貼。"),
        ("6. 既有通知表一鍵還原與併排修正：", "支援載入既有運輸通知表還原原始列位，填寫修正到廠時間後自動於右側併排產生紅字加粗與刪除線之出貨修正卡片。")
    ]

    for b_title, b_desc in bullets:
        pb = doc.add_paragraph(style='List Bullet')
        pb.paragraph_format.space_after = Pt(4)
        rb1 = pb.add_run(b_title)
        rb1.font.name = 'Microsoft JhengHei'
        rb1.font.bold = True
        rb1.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
        rb2 = pb.add_run(b_desc)
        rb2.font.name = 'Microsoft JhengHei'

    # Section 2 Header
    p6 = doc.add_paragraph()
    p6.paragraph_format.space_before = Pt(14)
    p6.paragraph_format.space_after = Pt(6)
    r6 = p6.add_run("二、 圖文流程操作教學 (Step-by-Step)")
    r6.font.name = 'Microsoft JhengHei'
    r6.font.size = Pt(15)
    r6.font.bold = True
    r6.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    steps = [
        (
            "【步驟 1】對照表檔案狀態與免重開一鍵刷新",
            "軟體開啟時會自動核對範本與地點對照表。若在對照表中新增了地點，無需關閉重開程式，直接點擊「重新載入對照表」按鈕即可即時刷新對照碼！",
            "image1.png",
            "▲ 圖 1：系統狀態列與『重新載入對照表』按鈕位置示意圖"
        ),
        (
            "【步驟 2】智慧多欄 Ctrl+V 貼上與欄位自動歸位",
            "直接從 Excel 複製多筆資料，於「批號」輸入框按下 Ctrl+V。無論複製順序為【到貨/批號/槽號/地點】或【批號/品名/地點/日期/時間】，系統均能智慧精準歸位！",
            "image2.png",
            "▲ 圖 2：智慧多欄複製貼上自動歸位示意圖"
        ),
        (
            "【步驟 3】一鍵全清與單列獨立清空",
            "填錯資料時無需逐一退格刪除！單列填錯點擊右側紅色「清空」按鈕；整體重填點擊右上角「清除全部資料 (全清)」按鈕即可秒級復原。",
            "image3.png",
            "▲ 圖 3：一鍵全清與單列清空按鈕位置示意圖"
        ),
        (
            "【步驟 4】上傳 COA 檢驗截圖與自動 OCR 辨識貼入",
            "點擊頂部橘色「上傳 COA 截圖」按鈕，可一次選取多張當日檢驗截圖檔。產出三合一單時，系統自動透過 OCR 引擎解析對應批號，將檢驗表頭與數據自動裁切並精確嵌入三合一單 F5 儲存格！",
            "image_coa.png",
            "▲ 圖 4：頂部『上傳 COA 截圖』按鈕位置示意圖"
        ),
        (
            "【步驟 5】勝一訂單匯入：台積電專屬過濾與「當天至+2天」優先推薦",
            "點擊「從 Excel 匯入」選取『勝一訂單.xlsx』時，系統會自動過濾非台積電廠商，並優先鎖定「當天至+2天 (今天~後天)」的最新台積電排程自動預選！亦可透過快捷按鈕單獨切換「今天」、「明天」、「後天」或倒數筆數。",
            "image_shinyi_import.png",
            "▲ 圖 5：勝一訂單台積電專屬過濾與當天至+2天優先推薦示意圖"
        ),
        (
            "【步驟 6】還原既有通知表與點對點修正到廠時間",
            "當下午需要修正上午產生的通知表時，點擊「載入既有『運輸通知表』修訂」（預設開啟當天資料夾）。一鍵還原所有原始列號與 10~11 碼批號，只需在該列輸入「修正到廠時間」即可產出併排修正卡片！",
            "image5.png",
            "▲ 圖 6：載入既有通知表與填寫修正時間示意圖"
        ),
        (
            "【步驟 7】一鍵產生報表與全新【槽號】檔名格式",
            "勾選欲產生的報表種類後，點擊下方「開始批次產生 Excel 報表」。三合一單最新檔名規格加入【槽號】，徹底解決同天同地點出車檔案覆蓋的問題！",
            "image6.png",
            "▲ 圖 7：一鍵批次產生報表與槽號檔名格式示意圖"
        )
    ]

    for title, desc, img_name, caption in steps:
        # Step Header
        ps_h = doc.add_paragraph()
        ps_h.paragraph_format.space_before = Pt(10)
        ps_h.paragraph_format.space_after = Pt(4)
        rs_h = ps_h.add_run(title)
        rs_h.font.name = 'Microsoft JhengHei'
        rs_h.font.size = Pt(12.5)
        rs_h.font.bold = True
        rs_h.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

        # Step Desc
        ps_d = doc.add_paragraph()
        rs_d = ps_d.add_run(desc)
        rs_d.font.name = 'Microsoft JhengHei'

        # Step Image
        img_path = os.path.join(assets_dir, img_name)
        if not os.path.exists(img_path):
            img_path = os.path.join(base_dir, img_name)
        
        if os.path.exists(img_path):
            ps_i = doc.add_paragraph()
            ps_i.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ps_i.paragraph_format.space_before = Pt(6)
            ps_i.paragraph_format.space_after = Pt(2)
            ps_i.add_run().add_picture(img_path, width=Inches(6.2))

        # Step Caption
        ps_c = doc.add_paragraph()
        ps_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ps_c.paragraph_format.space_after = Pt(10)
        rs_c = ps_c.add_run(caption)
        rs_c.font.name = 'Microsoft JhengHei'
        rs_c.font.size = Pt(9.5)
        rs_c.font.bold = True
        rs_c.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # Section 3 Header
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(14)
    p_t.paragraph_format.space_after = Pt(6)
    r_t = p_t.add_run("三、 報表檔名與規格說明")
    r_t.font.name = 'Microsoft JhengHei'
    r_t.font.size = Pt(15)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    # Table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr_cells = table.rows[0].cells
    hdr_titles = ["報表類型", "檔名命名規格", "說明 / 特色"]
    hdr_widths = [Inches(1.5), Inches(3.2), Inches(2.3)]
    for i, title in enumerate(hdr_titles):
        cell = hdr_cells[i]
        cell.width = hdr_widths[i]
        set_cell_background(cell, "002060")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.font.name = 'Microsoft JhengHei'
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rows_data = [
        ("獨立三合一單", "[出貨日期]. [地點]_[槽號]_台積電槽車barcode三合一單.xlsx", "例如：2026.8.28. 15P7_S405_台積電槽車barcode三合一單.xlsx（含自動 QR Code 與 COA 檢驗截圖嵌入）"),
        ("運輸通知表", "運輸通知表.xlsx", "卡片標題與左側品名依產品動態調整（如 Shiny SEP73E5 Lorry），充填數量自動帶入固定重量；若有修正時間自動併排呈現修正卡片"),
        ("歸檔資料夾", "勝一三合一單輸出_YYYYMMDD", "例如：勝一三合一單輸出_20260828，當天所有產出集中於獨立資料夾")
    ]

    for row_idx, data in enumerate(rows_data):
        row_cells = table.add_row().cells
        bg_color = "F9F9F9" if row_idx % 2 == 1 else "FFFFFF"
        for i, text in enumerate(data):
            cell = row_cells[i]
            cell.width = hdr_widths[i]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(text)
                r.font.name = 'Microsoft JhengHei'
                r.font.bold = True
            else:
                r = p.add_run(text)
                r.font.name = 'Microsoft JhengHei'

    # Section 4 Header: Troubleshooting & FAQ
    p_faq = doc.add_paragraph()
    p_faq.paragraph_format.space_before = Pt(16)
    p_faq.paragraph_format.space_after = Pt(6)
    r_faq = p_faq.add_run("四、 常見問題與故障排除 (Troubleshooting & FAQ)")
    r_faq.font.name = 'Microsoft JhengHei'
    r_faq.font.size = Pt(15)
    r_faq.font.bold = True
    r_faq.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    faqs = [
        (
            "Q1：勝一訂單中有許多非台積電客戶（如南亞、長春），匯入時是否會混在一起？",
            "【解答】不會！系統在解析勝一訂單時會自動以台積電指送地與對象簡稱進行精準過濾，只保留純台積電出貨排程，確保操作畫面簡潔專注。"
        ),
        (
            "Q2：匯入勝一訂單時，指送地（如「台積12廠P1」）是否能自動辨識？",
            "【解答】可以！系統內建智慧地點正規化演算法，會自動將「台積12廠P1」清理為「12P1」、「台積新竹14廠-P5」轉換為「14P5」、「台積先進AP8」轉換為「AP8」，並自動匹配對照表填入長代號。"
        ),
        (
            "Q3：勝一訂單批號為 11 碼時（如 26707S125T1），槽號是否能正常提取？",
            "【解答】可以！系統已升級支援 10~11 碼勝一批號規則，會精準擷取中間的英數槽號（如 S125、S96、S405、E305、E319），絕不報錯。"
        ),
        (
            "Q4：運輸通知單上的品名與充填數量如何運作？",
            "【解答】品名會自動由勝一訂單的品名欄位帶出；充填數量由系統內建之「產品-固定重量對照表」自動填入（如 4300 KG），日後各產品重量若有固定指定值，可直接設定至對照表，產出時全自動填入，不需現場人員手動逐筆修改。"
        ),
        (
            "Q5：地點輸入後「長代號」欄位顯示紅色「對照表中找不到」？",
            "【原因與排除】表示輸入之地點代號尚未登錄在『地點代號對照表.xlsx』中。排除步驟：\n"
            "1. 開啟軟體同目錄下的『地點代號對照表.xlsx』檔案。\n"
            "2. 在下方新增對應之地點名稱與長代號後儲存檔案。\n"
            "3. 回到軟體視窗點擊頂部「重新載入對照表」按鈕即可立即生效，無需關閉重啟程式！"
        ),
        (
            "Q6：COA 檢驗截圖上傳後，產出之三合一單 F5 儲存格未帶出檢驗圖？",
            "【原因與排除】請確認：\n"
            "1. COA 截圖內的「批號」文字請保持清晰、勿過度壓縮或遮蔽。\n"
            "2. 若系統提示尚未安裝 Tesseract-OCR，請依引導下載安裝於預設路徑即可自動啟用。"
        ),
        (
            "Q7：點擊「開始批次產生 Excel 報表」時跳出「Permission Denied (權限遭拒)」？",
            "【原因與排除】同名 Excel 檔案目前正被人員以 Excel 開啟檢視中導致鎖檔。請先將開啟中的 Excel 關閉後，再次點擊產生按鈕即可。"
        )
    ]

    for q_title, a_desc in faqs:
        pq = doc.add_paragraph()
        pq.paragraph_format.space_before = Pt(8)
        pq.paragraph_format.space_after = Pt(2)
        rq = pq.add_run(q_title)
        rq.font.name = 'Microsoft JhengHei'
        rq.font.size = Pt(11.5)
        rq.font.bold = True
        rq.font.color.rgb = RGBColor(0x00, 0x50, 0x9E)

        pa = doc.add_paragraph()
        pa.paragraph_format.space_before = Pt(0)
        pa.paragraph_format.space_after = Pt(6)
        ra = pa.add_run(a_desc)
        ra.font.name = 'Microsoft JhengHei'
        ra.font.size = Pt(10.5)

    doc.save(docx_path)
    print(f"Word file generated: {docx_path}")

def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        word.Quit()
        print(f"PDF file generated: {pdf_path}")
    except Exception as e:
        print(f"Failed to convert to PDF via Word COM: {e}")

if __name__ == "__main__":
    docx_file = os.path.join(base_dir, "勝一三合一單產生系統_操作手冊.docx")
    pdf_file = os.path.join(base_dir, "勝一三合一單產生系統_操作手冊.pdf")
    
    create_manual_docx(docx_file)
    convert_docx_to_pdf(docx_file, pdf_file)
