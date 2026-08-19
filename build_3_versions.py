import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import win32com.client

img_factory = r"C:\Users\C606-PC\.gemini\antigravity\brain\44faa0bd-7fa9-4497-a578-37048fb486f2\.user_uploaded\media_1787098103540.jpg"
img_office = r"C:\Users\C606-PC\.gemini\antigravity\brain\44faa0bd-7fa9-4497-a578-37048fb486f2\.user_uploaded\media_1787098101697.jpg"
base_dir = r"d:\GOOGLE ANGET"

steps_data = [
    ("步驟一：即時保護與隔離", "開啟有效糾正與補救措施", "接獲申訴或知悉情事時，雇主應立即啟動保護機制，調整排班座位或適當職務隔離，避免再次受害或二次報復。"),
    ("步驟二：獨立調查與申訴處理", "啟動獨立調查程序與性平委員會", "成立獨立申訴調查小組（外聘專家學者或性平委員比例應符合法定標準），給予雙方充分陳述意見之機會，全程客觀保密。"),
    ("步驟三：決議懲處與個案關懷", "作成調查報告、懲處與關懷資源", "依調查結果進行事實認定與權責懲處（申誡、記過、解僱等），並提供申訴人心理諮商、醫療輔導與法律協助，依規通報。"),
    ("步驟四：宣導教育與滾動追蹤", "全員定期宣導與組織文化維護", "針對調查發現之組織漏洞進行宣導，定期實施性騷擾與霸凌防治教育訓練，並持續追蹤當事人工作狀況，杜絕排擠。")
]

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def apply_font(run, size=10.5, bold=False, color_rgb=(0x33, 0x33, 0x33), italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Microsoft JhengHei" w:hAnsi="Microsoft JhengHei" w:eastAsia="Microsoft JhengHei" w:cs="Microsoft JhengHei"/>')
    rPr.append(rFonts)

# ==========================================
# 版本一：橫向大字簡報風格 (Landscape Large Font - 10.5pt ~ 11pt)
# ==========================================
def build_version_1(docx_path):
    doc = docx.Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # 標題 (放大至 22pt)
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_t.add_run("【職場霸凌與性騷擾案例分析與法規辨識手冊】")
    apply_font(r_t, size=22, bold=True, color_rgb=(0x00, 0x20, 0x60))
    p_t.paragraph_format.space_after = Pt(4)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s = p_sub.add_run("廠區作業與辦公室情境實務案例彙整評析 (大字簡報宣導版)")
    apply_font(r_s, size=12, bold=True, color_rgb=(0xC0, 0x00, 0x00))
    p_sub.paragraph_format.space_after = Pt(10)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        apply_font(r, size=15, bold=True, color_rgb=(0x00, 0x20, 0x60))
        return p

    # 一、廠區案例
    add_h1("一、 廠區作業現場情境案例評析")
    if os.path.exists(img_factory):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(8)
        r = p_img.add_run()
        r.add_picture(img_factory, width=Inches(9.6))

    table1 = doc.add_table(rows=1, cols=6)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = False
    
    hdr1 = table1.rows[0].cells
    hdr1_titles = ["案例類型", "發生場景", "行為樣態", "霸凌成分", "性騷擾成分", "加害者常見藉口"]
    hdr1_widths = [Inches(1.6), Inches(1.4), Inches(2.8), Inches(2.1), Inches(2.1), Inches(1.6)]
    
    for i, title in enumerate(hdr1_titles):
        cell = hdr1[i]
        cell.width = hdr1_widths[i]
        set_cell_background(cell, "002060")
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        apply_font(r, size=11, bold=True, color_rgb=(0xFF, 0xFF, 0xFF))

    data1 = [
        ("【案例一】\n假借「玩笑」肢體侵犯", "更衣室、狹窄作業區", "資深員工以「迎新、好玩」為由，對新進男員工進行「阿魯巴」、強行拉扯衣物，或碰觸臀部與下體。", "利用資歷與人數優勢（多對一）使新人不敢反抗，形成權力霸凌。", "違反意願的私密部位觸碰，構成「肢體性騷擾」。", "「大家都是男人，摸一下會少一塊肉嗎？」"),
        ("【案例二】\n針對「性別氣質」言語羞辱", "交接班會議、搬運現場", "領班針對體能較弱或性格溫和的男員工，公開辱罵「娘砲」、「沒卵葩」等貶抑字眼，並分配極重工作。", "聯合其他同事孤立受害者，並濫用職權給予不當勞動條件。", "針對性別特徵與氣質進行貶抑，構成「性別氣質騷擾」。", "「做粗活還這麼嬌弱，我是為了鍛鍊他。」"),
        ("【案例三】\n權力展現「敵意環境」", "吸菸區、現場通訊群組", "小主管常在休息時間強迫展示色情影片，當眾詢問性生活細節。反感者會在排班考績上遭到刁難。", "利用排班與考評權力進行威脅與報復，屬於職權霸凌。", "散布色情圖文並製造冒犯恐懼氛圍，構成「敵意環境性騷擾」。", "「開個黃腔而已，也太開不起玩笑了。」")
    ]

    for r_idx, row in enumerate(data1):
        row_cells = table1.add_row().cells
        bg_col = "F9F9F9" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = row_cells[c_idx]
            cell.width = hdr1_widths[c_idx]
            set_cell_background(cell, bg_col)
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                apply_font(r, size=10.5, bold=True, color_rgb=(0x00, 0x20, 0x60))
            elif c_idx == 5:
                apply_font(r, size=10, italic=True, bold=True, color_rgb=(0xC0, 0x00, 0x00))
            else:
                apply_font(r, size=10, color_rgb=(0x33, 0x33, 0x33))

    doc.add_page_break()

    # 二、辦公室案例
    add_h1("二、 辦公室與社交情境案例評析")
    if os.path.exists(img_office):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(8)
        r = p_img.add_run()
        r.add_picture(img_office, width=Inches(9.6))

    table2 = doc.add_table(rows=1, cols=6)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.autofit = False
    
    hdr2 = table2.rows[0].cells
    for i, title in enumerate(hdr1_titles):
        cell = hdr2[i]
        cell.width = hdr1_widths[i]
        set_cell_background(cell, "1B5E20")
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        apply_font(r, size=11, bold=True, color_rgb=(0xFF, 0xFF, 0xFF))

    data2 = [
        ("【案例一】\n通訊軟體與言語黃色玩笑", "茶水間、部門 LINE 群組", "同事常在群組傳送色情梗圖，或在茶水間公開拿男同事性生活或「斯文氣質」開黃腔嘲笑。", "透過公開嘲弄建立社交優勢地位，屬於言語霸凌。", "針對性特徵與性別氣質嘲笑並散布圖文，構成「言語與視覺性騷擾」。", "「這只是男生之間的幹話，幹嘛這麼嚴肅？」"),
        ("【案例二】\n假借放鬆的「肢體越界」", "辦公座位區、員工休息區", "以「看你最近壓力很大」為由強行幫男同事「按摩」肩頸，或走道交錯時拍打對方臀部、大腿。", "無視他人身體界線，利用體型強迫接受接觸，屬於行為霸凌。", "違反當事人意願，觸碰具有性意涵部位，構成「肢體性騷擾」。", "「大家都是兄弟，抓一下肩膀又不會懷孕。」"),
        ("【案例三】\n「兄弟幫」社交排擠針對", "跨部門會議、應酬聚會", "辦公室形成特定「兄弟小團體」。不參與粗俗玩笑者，專案被隱瞞資訊或會議上被貶低意見。", "刻意孤立特定對象並阻礙執行工作，屬於關係霸凌與排擠。", "營造「不加入開黃腔就不合群」氛圍，構成「敵意環境性騷擾」。", "「他自己不合群太難相處，我們才不想跟他合作。」")
    ]

    for r_idx, row in enumerate(data2):
        row_cells = table2.add_row().cells
        bg_col = "F9F9F9" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = row_cells[c_idx]
            cell.width = hdr1_widths[c_idx]
            set_cell_background(cell, bg_col)
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                apply_font(r, size=10.5, bold=True, color_rgb=(0x1B, 0x5E, 0x20))
            elif c_idx == 5:
                apply_font(r, size=10, italic=True, bold=True, color_rgb=(0xC0, 0x00, 0x00))
            else:
                apply_font(r, size=10, color_rgb=(0x33, 0x33, 0x33))

    doc.add_page_break()

    # 三、法規與處置四步驟
    add_h1("三、 關鍵法規辨識與企業處置四步驟")
    
    table_steps = doc.add_table(rows=1, cols=3)
    table_steps.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_steps.autofit = False
    
    hdr_s = table_steps.rows[0].cells
    hdr_s_titles = ["處置步驟", "核心作業項目", "具體執行內容與法規要求"]
    hdr_s_widths = [Inches(2.2), Inches(2.8), Inches(6.6)]
    
    for i, title in enumerate(hdr_s_titles):
        cell = hdr_s[i]
        cell.width = hdr_s_widths[i]
        set_cell_background(cell, "002060")
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        apply_font(r, size=11, bold=True, color_rgb=(0xFF, 0xFF, 0xFF))

    for r_idx, (s_num, s_title, s_desc) in enumerate(steps_data):
        row_cells = table_steps.add_row().cells
        bg_col = "F9F9F9" if r_idx % 2 == 1 else "FFFFFF"
        
        cell0 = row_cells[0]
        cell0.width = hdr_s_widths[0]
        set_cell_background(cell0, bg_col)
        set_cell_margins(cell0, top=80, bottom=80, left=80, right=80)
        p0 = cell0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(s_num)
        apply_font(r0, size=10.5, bold=True, color_rgb=(0xC0, 0x00, 0x00))
        
        cell1 = row_cells[1]
        cell1.width = hdr_s_widths[1]
        set_cell_background(cell1, bg_col)
        set_cell_margins(cell1, top=80, bottom=80, left=80, right=80)
        p1 = cell1.paragraphs[0]
        r1 = p1.add_run(s_title)
        apply_font(r1, size=10.5, bold=True, color_rgb=(0x00, 0x20, 0x60))
        
        cell2 = row_cells[2]
        cell2.width = hdr_s_widths[2]
        set_cell_background(cell2, bg_col)
        set_cell_margins(cell2, top=80, bottom=80, left=80, right=80)
        p2 = cell2.paragraphs[0]
        r2 = p2.add_run(s_desc)
        apply_font(r2, size=10, color_rgb=(0x33, 0x33, 0x33))

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 警示框
    tbl_card = doc.add_table(rows=1, cols=1)
    tbl_card.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_card = tbl_card.cell(0, 0)
    set_cell_background(c_card, "FFF8E7")
    set_cell_margins(c_card, top=100, bottom=100, left=120, right=120)
    p_c = c_card.paragraphs[0]
    r_ct = p_c.add_run("【企業法定義務提醒】")
    apply_font(r_ct, size=11, bold=True, color_rgb=(0xC0, 0x00, 0x00))
    r_cd = p_c.add_run("當接獲申訴或知悉性騷擾/霸凌情事時，雇主必須立即採取「有效之糾正及補救措施」（包含保護申訴人、啟動獨立調查程序、隔離加害者、給予心理諮商支援），否則依法最高可處新臺幣 100 萬元罰鍰。")
    apply_font(r_cd, size=10.5, color_rgb=(0x33, 0x33, 0x33))

    doc.save(docx_path)
    print(f"Version 1 (Landscape Large Font) saved to: {docx_path}")

# ==========================================
# 版本二：直向大字清晰版 (Portrait Large Font - 10pt~10.5pt)
# ==========================================
def build_version_2(docx_path):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        apply_font(r, size=14, bold=True, color_rgb=(0x00, 0x20, 0x60))
        return p

    # --- PAGE 1 ---
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_t.add_run("【職場霸凌與性騷擾案例分析與法規辨識手冊】")
    apply_font(r_t, size=18, bold=True, color_rgb=(0x00, 0x20, 0x60))
    p_t.paragraph_format.space_after = Pt(2)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s = p_sub.add_run("廠區作業與辦公室情境實務案例彙整評析 (直向大字清晰版)")
    apply_font(r_s, size=10.5, bold=True, color_rgb=(0xC0, 0x00, 0x00))
    p_sub.paragraph_format.space_after = Pt(6)

    add_h1("一、 廠區作業現場情境案例評析")
    if os.path.exists(img_factory):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(6)
        r = p_img.add_run()
        r.add_picture(img_factory, width=Inches(6.8))

    table1 = doc.add_table(rows=1, cols=6)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = False
    
    hdr1 = table1.rows[0].cells
    hdr1_titles = ["案例類型", "發生場景", "行為樣態", "霸凌成分", "性騷擾成分", "加害者常見藉口"]
    hdr1_widths = [Inches(1.2), Inches(1.1), Inches(1.8), Inches(1.4), Inches(1.4), Inches(1.1)]
    
    for i, title in enumerate(hdr1_titles):
        cell = hdr1[i]
        cell.width = hdr1_widths[i]
        set_cell_background(cell, "002060")
        set_cell_margins(cell, top=60, bottom=60, left=50, right=50)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        apply_font(r, size=10, bold=True, color_rgb=(0xFF, 0xFF, 0xFF))

    data1 = [
        ("【案例一】\n假借「玩笑」肢體侵犯", "更衣室、狹窄作業區", "資深員工以迎新好玩為由對新進男員工拉扯衣物或碰觸臀部下體。", "利用資歷人數優勢（多對一）使新人不敢反抗，形成權力霸凌。", "違反意願私密部位觸碰，構成肢體性騷擾。", "「大家都是男人，摸一下會少一塊肉嗎？」"),
        ("【案例二】\n針對「性別氣質」言語羞辱", "交接班會議、搬運現場", "領班針對體能較弱者辱罵「娘砲」、「沒卵葩」並分配極重工作。", "聯合同事孤立受害者，濫用職權給予不當勞動條件。", "針對性別特徵與氣質貶抑，構成性別氣質騷擾。", "「做粗活還這麼嬌弱，我是為了鍛鍊他。」"),
        ("【案例三】\n權力展現「敵意環境」", "吸菸區、現場通訊群組", "小主管常強迫展示色情影片、詢問性生活細節，反感者考績遭刁難。", "利用排班考評權力威脅報復，屬於職權霸凌。", "散布色情圖文並製造冒犯恐懼，構成敵意環境性騷擾。", "「開個黃腔而已，也太開不起玩笑了。」")
    ]

    for r_idx, row in enumerate(data1):
        row_cells = table1.add_row().cells
        bg_col = "F9F9F9" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = row_cells[c_idx]
            cell.width = hdr1_widths[c_idx]
            set_cell_background(cell, bg_col)
            set_cell_margins(cell, top=60, bottom=60, left=50, right=50)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                apply_font(r, size=9.5, bold=True, color_rgb=(0x00, 0x20, 0x60))
            elif c_idx == 5:
                apply_font(r, size=9, italic=True, bold=True, color_rgb=(0xC0, 0x00, 0x00))
            else:
                apply_font(r, size=9.5, color_rgb=(0x33, 0x33, 0x33))

    doc.add_page_break()

    # --- PAGE 2 ---
    add_h1("二、 辦公室與社交情境案例評析")
    if os.path.exists(img_office):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(6)
        r = p_img.add_run()
        r.add_picture(img_office, width=Inches(6.8))

    table2 = doc.add_table(rows=1, cols=6)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.autofit = False
    
    hdr2 = table2.rows[0].cells
    for i, title in enumerate(hdr1_titles):
        cell = hdr2[i]
        cell.width = hdr1_widths[i]
        set_cell_background(cell, "1B5E20")
        set_cell_margins(cell, top=60, bottom=60, left=50, right=50)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        apply_font(r, size=10, bold=True, color_rgb=(0xFF, 0xFF, 0xFF))

    data2 = [
        ("【案例一】\n通訊軟體與言語黃色玩笑", "茶水間、部門 LINE 群組", "同事常在群組傳色情梗圖，或在茶水間拿男同事性生活與斯文氣質開黃腔。", "透過公開嘲弄建立社交優勢地位，屬於言語霸凌。", "針對性特徵與氣質嘲笑並散布圖文，構成言語與視覺性騷擾。", "「這只是男生之間的幹話，幹嘛這麼嚴肅？」"),
        ("【案例二】\n假借放鬆的「肢體越界」", "辦公座位區、員工休息區", "以壓力大為由強行幫男同事「按摩」肩頸，或走道交錯時拍打對方臀部大腿。", "無視他人身體界線，利用體型強迫接受接觸，屬行為霸凌。", "違反意願觸碰具性意涵部位，構成肢體性騷擾。", "「大家都是兄弟，抓一下肩膀又不會懷孕。」"),
        ("【案例三】\n「兄弟幫」社交排擠與針對", "跨部門會議、應酬聚會", "辦公室形成兄弟小團體。不參與粗俗玩笑者專案被隱瞞資訊或會議被貶低。", "刻意孤立特定對象並阻礙執行工作，屬於關係霸凌排擠。", "營造「不加入開黃腔就不合群」氛圍，構成敵意環境性騷擾。", "「他自己不合群太難相處，我們才不想跟他合作。」")
    ]

    for r_idx, row in enumerate(data2):
        row_cells = table2.add_row().cells
        bg_col = "F9F9F9" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = row_cells[c_idx]
            cell.width = hdr1_widths[c_idx]
            set_cell_background(cell, bg_col)
            set_cell_margins(cell, top=60, bottom=60, left=50, right=50)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                apply_font(r, size=9.5, bold=True, color_rgb=(0x1B, 0x5E, 0x20))
            elif c_idx == 5:
                apply_font(r, size=9, italic=True, bold=True, color_rgb=(0xC0, 0x00, 0x00))
            else:
                apply_font(r, size=9.5, color_rgb=(0x33, 0x33, 0x33))

    doc.add_page_break()

    # --- PAGE 3 ---
    add_h1("三、 關鍵法規辨識與企業處置四步驟")
    
    p_b1 = doc.add_paragraph(style='List Bullet')
    p_b1.paragraph_format.space_after = Pt(2)
    r = p_b1.add_run("• 意願優先原則：")
    apply_font(r, size=10, bold=True, color_rgb=(0x00, 0x20, 0x60))
    r2 = p_b1.add_run("性騷擾成立不限於異性之間，只要違反當事人意願並感到冒犯即成立。")
    apply_font(r2, size=10)

    p_b2 = doc.add_paragraph(style='List Bullet')
    p_b2.paragraph_format.space_after = Pt(6)
    r = p_b2.add_run("• 性別氣質保護與敵意環境：")
    apply_font(r, size=10, bold=True, color_rgb=(0x00, 0x20, 0x60))
    r2 = p_b2.add_run("針對「娘砲」、「太斯文」羞辱或散布色情梗圖建立排擠文化，皆屬違法敵意環境。")
    apply_font(r2, size=10)

    table_steps = doc.add_table(rows=1, cols=3)
    table_steps.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_steps.autofit = False
    
    hdr_s = table_steps.rows[0].cells
    hdr_s_titles = ["處置步驟", "核心作業項目", "具體執行內容與法規要求"]
    hdr_s_widths = [Inches(1.5), Inches(1.8), Inches(4.7)]
    
    for i, title in enumerate(hdr_s_titles):
        cell = hdr_s[i]
        cell.width = hdr_s_widths[i]
        set_cell_background(cell, "002060")
        set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        apply_font(r, size=10, bold=True, color_rgb=(0xFF, 0xFF, 0xFF))

    for r_idx, (s_num, s_title, s_desc) in enumerate(steps_data):
        row_cells = table_steps.add_row().cells
        bg_col = "F9F9F9" if r_idx % 2 == 1 else "FFFFFF"
        
        cell0 = row_cells[0]
        cell0.width = hdr_s_widths[0]
        set_cell_background(cell0, bg_col)
        set_cell_margins(cell0, top=60, bottom=60, left=60, right=60)
        p0 = cell0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(s_num)
        apply_font(r0, size=9.5, bold=True, color_rgb=(0xC0, 0x00, 0x00))
        
        cell1 = row_cells[1]
        cell1.width = hdr_s_widths[1]
        set_cell_background(cell1, bg_col)
        set_cell_margins(cell1, top=60, bottom=60, left=60, right=60)
        p1 = cell1.paragraphs[0]
        r1 = p1.add_run(s_title)
        apply_font(r1, size=9.5, bold=True, color_rgb=(0x00, 0x20, 0x60))
        
        cell2 = row_cells[2]
        cell2.width = hdr_s_widths[2]
        set_cell_background(cell2, bg_col)
        set_cell_margins(cell2, top=60, bottom=60, left=60, right=60)
        p2 = cell2.paragraphs[0]
        r2 = p2.add_run(s_desc)
        apply_font(r2, size=9.5, color_rgb=(0x33, 0x33, 0x33))

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    tbl_card = doc.add_table(rows=1, cols=1)
    tbl_card.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_card = tbl_card.cell(0, 0)
    set_cell_background(c_card, "FFF8E7")
    set_cell_margins(c_card, top=80, bottom=80, left=100, right=100)
    p_c = c_card.paragraphs[0]
    r_ct = p_c.add_run("【企業法定義務提醒】")
    apply_font(r_ct, size=10, bold=True, color_rgb=(0xC0, 0x00, 0x00))
    r_cd = p_c.add_run("當接獲申訴或知悉性騷擾/霸凌情事時，雇主必須立即採取「有效之糾正及補救措施」（包含保護申訴人、啟動獨立調查程序、隔離加害者、給予心理諮商支援），否則依法最高可處新臺幣 100 萬元罰鍰。")
    apply_font(r_cd, size=9.5, color_rgb=(0x33, 0x33, 0x33))

    doc.save(docx_path)
    print(f"Version 2 (Portrait Large Font) saved to: {docx_path}")

# ==========================================
# 版本三：大字圖文卡片風格 (Card Modules Zero Emoji - 11pt)
# ==========================================
def build_version_3(docx_path):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        apply_font(r, size=15, bold=True, color_rgb=(0x00, 0x20, 0x60))
        return p

    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_t.add_run("【職場霸凌與性騷擾案例分析與法規辨識手冊】")
    apply_font(r_t, size=20, bold=True, color_rgb=(0x00, 0x20, 0x60))
    p_t.paragraph_format.space_after = Pt(2)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s = p_sub.add_run("圖文資訊大字卡片版 (男同儕與職場性別友善專題)")
    apply_font(r_s, size=11, bold=True, color_rgb=(0xC0, 0x00, 0x00))
    p_sub.paragraph_format.space_after = Pt(10)

    # 一、廠區卡片
    add_h1("一、 廠區作業現場情境案例卡片")
    if os.path.exists(img_factory):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(8)
        r = p_img.add_run()
        r.add_picture(img_factory, width=Inches(6.8))

    cards_factory = [
        ("【案例一】假借「玩笑」的肢體侵犯", "更衣室、狹窄作業區", "資深員工以迎新好玩為由對新進男員工拉扯衣物或碰觸臀部下體。", "多對一資歷壓迫（權力霸凌）", "私密部位違意觸碰（肢體性騷擾）", "「大家都是男人，摸一下會少一塊肉嗎？」"),
        ("【案例二】針對「性別氣質」的言語羞辱", "交接班會議、搬運現場", "領班針對體能較弱者辱罵「娘砲」、「沒卵葩」並分配極重工作。", "聯合孤立與給予不當條件（針對性霸凌）", "性別特徵與氣質貶抑（性別氣質騷擾）", "「做粗活還這麼嬌弱，我是為了鍛鍊他。」"),
        ("【案例三】權力展現的「敵意環境」建立", "吸菸區、現場通訊群組", "小主管常強迫展示色情影片、詢問性生活細節，反感者考績遭刁難。", "排班考評威脅報復（職權霸凌）", "色情圖文與冒犯恐懼氛圍（敵意環境性騷擾）", "「開個黃腔而已，也太開不起玩笑了。」")
    ]

    for title, loc, act, b_comp, s_comp, excuse in cards_factory:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F0F4F8")
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        r_t = p.add_run(f"{title} （發生場景：{loc}）\n")
        apply_font(r_t, size=11.5, bold=True, color_rgb=(0x00, 0x20, 0x60))
        
        r_a = p.add_run(f"• 行為樣態：{act}\n")
        apply_font(r_a, size=10.5, color_rgb=(0x33, 0x33, 0x33))
        
        r_b = p.add_run(f"• 霸凌成分：{b_comp}  |  • 性騷擾成分：{s_comp}\n")
        apply_font(r_b, size=10.5, bold=True, color_rgb=(0x1B, 0x5E, 0x20))
        
        r_e = p.add_run(f"◆ 加害常見藉口：{excuse}")
        apply_font(r_e, size=10, italic=True, bold=True, color_rgb=(0xC0, 0x00, 0x00))
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # 二、辦公室卡片
    add_h1("二、 辦公室與社交情境案例卡片")
    if os.path.exists(img_office):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(8)
        r = p_img.add_run()
        r.add_picture(img_office, width=Inches(6.8))

    cards_office = [
        ("【案例一】通訊軟體與言語的「黃色玩笑」", "茶水間、部門 LINE 群組", "同事常在群組傳色情梗圖，或在茶水間拿男同事性生活與斯文氣質開黃腔。", "公開嘲弄建立優勢（言語霸凌）", "針對氣質嘲笑與散布圖文（言語與視覺性騷擾）", "「這只是男生之間的幹話，幹嘛這麼嚴肅？」"),
        ("【案例二】假借放鬆的「肢體越界」", "辦公座位區、員工休息區", "以壓力大為由強行幫男同事「按摩」肩頸，或走道交錯時拍打對方臀部大腿。", "強迫接受接觸（行為霸凌）", "違反意願觸碰性意涵部位（肢體性騷擾）", "「大家都是兄弟，抓一下肩膀又不會懷孕。」"),
        ("【案例三】「兄弟幫」社交排擠與針對", "跨部門會議、應酬聚會", "辦公室形成兄弟小團體。不參與粗俗玩笑者專案被隱瞞資訊或會議被貶低。", "刻意孤立與阻礙工作（關係霸凌與排擠）", "營造不合群氛圍（敵意環境性騷擾）", "「他自己不合群太難相處，我們才不想跟他合作。」")
    ]

    for title, loc, act, b_comp, s_comp, excuse in cards_office:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F0F8F0")
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        r_t = p.add_run(f"{title} （發生場景：{loc}）\n")
        apply_font(r_t, size=11.5, bold=True, color_rgb=(0x1B, 0x5E, 0x20))
        
        r_a = p.add_run(f"• 行為樣態：{act}\n")
        apply_font(r_a, size=10.5, color_rgb=(0x33, 0x33, 0x33))
        
        r_b = p.add_run(f"• 霸凌成分：{b_comp}  |  • 性騷擾成分：{s_comp}\n")
        apply_font(r_b, size=10.5, bold=True, color_rgb=(0x00, 0x20, 0x60))
        
        r_e = p.add_run(f"◆ 加害常見藉口：{excuse}")
        apply_font(r_e, size=10, italic=True, bold=True, color_rgb=(0xC0, 0x00, 0x00))
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 三、法規處置四步驟
    add_h1("三、 企業處置四步驟卡片")
    for s_num, s_title, s_desc in steps_data:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "FFF5F5")
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        r_t = p.add_run(f"【{s_num}】—— {s_title}\n")
        apply_font(r_t, size=11, bold=True, color_rgb=(0xC0, 0x00, 0x00))
        r_d = p.add_run(f"說明：{s_desc}")
        apply_font(r_d, size=10.5, color_rgb=(0x33, 0x33, 0x33))
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    doc.save(docx_path)
    print(f"Version 3 (Card Modules Zero Emoji) saved to: {docx_path}")

def convert_to_pdf(docx_path, pdf_path):
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        word.Quit()
        print(f"PDF generated: {pdf_path}")
    except Exception as e:
        print(f"PDF error: {e}")

if __name__ == "__main__":
    main_docx = os.path.join(base_dir, "職場霸凌與性騷擾案例分析與法規辨識手冊.docx")
    main_pdf = os.path.join(base_dir, "職場霸凌與性騷擾案例分析與法規辨識手冊.pdf")

    v1_docx = os.path.join(base_dir, "職場霸凌與性騷擾案例分析與法規辨識手冊_橫向簡報版.docx")
    v1_pdf = os.path.join(base_dir, "職場霸凌與性騷擾案例分析與法規辨識手冊_橫向簡報版.pdf")
    
    v2_docx = os.path.join(base_dir, "職場霸凌與性騷擾案例分析與法規辨識手冊_緊湊直向3頁版.docx")
    v2_pdf = os.path.join(base_dir, "職場霸凌與性騷擾案例分析與法規辨識手冊_緊湊直向3頁版.pdf")

    v3_docx = os.path.join(base_dir, "職場霸凌與性騷擾案例分析與法規辨識手冊_圖文模組卡片版.docx")
    v3_pdf = os.path.join(base_dir, "職場霸凌與性騷擾案例分析與法規辨識手冊_圖文模組卡片版.pdf")

    build_version_1(main_docx)
    convert_to_pdf(main_docx, main_pdf)

    build_version_1(v1_docx)
    convert_to_pdf(v1_docx, v1_pdf)

    build_version_2(v2_docx)
    convert_to_pdf(v2_docx, v2_pdf)

    build_version_3(v3_docx)
    convert_to_pdf(v3_docx, v3_pdf)
