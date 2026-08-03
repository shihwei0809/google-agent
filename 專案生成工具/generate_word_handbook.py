import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

doc = docx.Document()

# Adjust margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Set base styles
normal_style = doc.styles['Normal']
normal_style.font.name = 'Microsoft JhengHei'
normal_style.font.size = Pt(11)
normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("📋 勝一化工【彰濱廠區】助理管理師\n全務招募面試與實務評量手冊")
title_run.font.size = Pt(18)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("涵蓋：進出貨過磅、T100 ERP 登打、包材批號維護與現場驗收/盤點應變能力測驗")
sub_run.font.size = Pt(10)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Section 1
h1 = doc.add_heading(level=1)
h1_run = h1.add_run("🎯 一、 招募測驗設計架構")
h1_run.font.size = Pt(14)
h1_run.font.bold = True
h1_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

p_flow = doc.add_paragraph()
p_flow.paragraph_format.space_after = Pt(8)
p_flow.add_run("本招募測驗流程設計為 ").font.color.rgb = RGBColor(0x33, 0x33, 0x33)
run_b = p_flow.add_run("45 分鐘精準篩選流程")
run_b.bold = True
p_flow.add_run("，確保應徵者同時具備 Office 數據處理能力、ERP 邏輯觀念與現場化學品倉儲應變能力：")

flows = [
    ("第一階段 (15 分鐘)", "Excel 實機上機測驗：過磅數據對帳、VLOOKUP 串接、樞紐分析表。"),
    ("第二階段 (10 分鐘)", "ERP 與倉儲現場情境簡答題：進出貨單據異常、包材批號 FIFO、過磅防呆。"),
    ("第三階段 (15 分鐘)", "行為面談與現場環境適應力：彰濱廠區現場作業意願與跨部門溝通能力。"),
    ("第四階段 (5 分鐘)", "綜合評分與錄取決策：填寫面試官綜合評分表。")
]

for stage, desc in flows:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{stage}：")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    p.add_run(desc)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Section 2
h2 = doc.add_heading(level=1)
h2_run = h2.add_run("📁 二、 Excel 實機上機測驗題目大綱與雙版本檔案")
h2_run.font.size = Pt(14)
h2_run.font.bold = True
h2_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

p_ex = doc.add_paragraph()
p_ex.add_run("配套實作檔案（已拆分為 雙版本 Excel 檔）：\n").font.color.rgb = RGBColor(0x33, 0x33, 0x33)

r1 = p_ex.add_run("1. 考生測驗用檔：")
r1.bold = True
p_ex.add_run("《勝一化工_彰濱廠區_助理管理師_實機測驗題庫_考生用.xlsx》（無解答）\n")

r2 = p_ex.add_run("2. 面試官解答版：")
r2.bold = True
p_ex.add_run("《勝一化工_彰濱廠區_助理管理師_實機測驗題庫_面試官解答版.xlsx》（含評分標準與標準公式）\n")

table1 = doc.add_table(rows=4, cols=4)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
table1.autofit = False

headers1 = ["測驗頁籤", "實測重點與題目描述", "配分", "考核指標"]
widths1 = [Inches(1.5), Inches(2.8), Inches(0.8), Inches(1.8)]

hdr_cells = table1.rows[0].cells
for i, name in enumerate(headers1):
    hdr_cells[i].text = name
    set_cell_background(hdr_cells[i], "1F4E78")
    set_cell_margins(hdr_cells[i])
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

t1_data = [
    ("02_過磅數據對帳", "計算 [淨重 = 毛重 - 空重]，並找出 2 筆過磅異常資料（如：毛重小於空重、數字輸入錯誤），填寫異常原因。", "40 分", "化工數據敏感度、邏輯對帳與抓錯細心度"),
    ("03_T100料號VLOOKUP", "利用 VLOOKUP 或 XLOOKUP 函數，將右側【T100料號主檔】對應的標準料號自動帶入驗收明細表。", "30 分", "系統資料串接能力、Office 函數熟練度"),
    ("04_包材領用統計", "依【包材領用明細表】，建立樞紐分析表 (Pivot Table) 統計各部門領用各包材的總數量。", "30 分", "數據報表彙整效率、樞紐分析表應用")
]

for row_idx, data in enumerate(t1_data, start=1):
    row_cells = table1.rows[row_idx].cells
    for col_idx, text in enumerate(data):
        row_cells[col_idx].text = text
        set_cell_margins(row_cells[col_idx])
        if row_idx % 2 == 1:
            set_cell_background(row_cells[col_idx], "F2F2F2")
        p = row_cells[col_idx].paragraphs[0]
        if col_idx in [0, 2]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for row in table1.rows:
    for idx, width in enumerate(widths1):
        row.cells[idx].width = width

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Section 3
h3 = doc.add_heading(level=1)
h3_run = h3.add_run("📝 三、 ERP 與倉儲現場實務情境問答題 (口試/筆試庫)")
h3_run.font.size = Pt(14)
h3_run.font.bold = True
h3_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

q_list = [
    ("Q1. 【T100 ERP 系統與單據過帳控管】",
     "供應商送來 50 加侖鐵桶 120 個，但 T100 系統開立的採購單 (PO) 數量只有 100 個。現場司機急著卸貨離場，要求你先簽收打單，你會如何處理？",
     "❌ 不合格回答：直接在 T100 系統中修改採購單數量為 120 個並過帳簽收。\n⭕ 優秀回答：現場先實收 100 個並依單過帳；多餘 20 個暫開代保管單或退回，同時立即通報主管與採購同仁，確認是否開立變更單（PO變更）後方可補刷入庫。展現嚴謹的內部控制觀念。"),
    
    ("Q2. 【包材批號維護與先進先出 (FIFO) 原則】",
     "當天產線急需領取「20L PE塑膠小桶」，倉庫內現有舊批號 LOT-A（保存期限剩 2 個月，位於架子深處）與新批號 LOT-B（保存期限剩 12 個月，位於最外側）。發料人員想拿新批號，你會如何應對？",
     "❌ 不合格回答：隨便發料人員拿哪一個，有拿夠數量就好。\n⭕ 優秀回答：堅持先進先出（FIFO）原則，要求優先領取舊批號 LOT-A。若擺放不易搬運，會記錄並協助改善倉儲位標示，避免化學品包材因積壓過期而造成公司損失。"),
    
    ("Q3. 【過磅異常與防呆機制】",
     "槽車進廠過磅時，過磅系統顯示空重與上週歷史紀錄相差 800 kg。司機解釋是「加滿了柴油和更換車頭」，你會直接進行毛重過磅並開單嗎？",
     "❌ 不合格回答：聽信司機口頭說明，直接扣除該空重並出貨開單。\n⭕ 優秀回答：過磅淨重直接影響化學品交易金額與品質。不會光憑口頭說詞，會要求查看車頭變更證明或執行二次覆磅確認，並通報主管。")
]

for title, q_text, a_text in q_list:
    qp = doc.add_paragraph()
    qp.paragraph_format.space_before = Pt(6)
    qp.paragraph_format.space_after = Pt(2)
    qr = qp.add_run(title)
    qr.bold = True
    qr.font.size = Pt(12)
    qr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    p_q = doc.add_paragraph()
    p_q.paragraph_format.left_indent = Inches(0.2)
    p_q.paragraph_format.space_after = Pt(4)
    r_qtitle = p_q.add_run("情境問題：")
    r_qtitle.bold = True
    p_q.add_run(q_text)
    
    for line in a_text.split('\n'):
        p_line = doc.add_paragraph()
        p_line.paragraph_format.left_indent = Inches(0.4)
        p_line.paragraph_format.space_after = Pt(2)
        if line.startswith("❌"):
            r = p_line.add_run(line)
            r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        else:
            r = p_line.add_run(line)
            r.font.color.rgb = RGBColor(0x38, 0x57, 0x23)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Section 4
h4 = doc.add_heading(level=1)
h4_run = h4.add_run("📊 四、 面試官綜合評分表 (Scorecard)")
h4_run.font.size = Pt(14)
h4_run.font.bold = True
h4_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

p_info = doc.add_paragraph()
p_info.paragraph_format.space_after = Pt(6)
p_info.add_run("應徵者姓名：____________________    面試日期：2026 年 _____ 月 _____ 日\n面試官簽名：____________________    綜合建議： [  ] 錄取   [  ] 備取   [  ] 不錄取").bold = True

table2 = doc.add_table(rows=6, cols=5)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.autofit = False

headers2 = ["評核項目", "權重", "評分標準 (1 - 5 分)", "得分", "面試官評語"]
widths2 = [Inches(1.6), Inches(0.8), Inches(2.6), Inches(0.7), Inches(1.8)]

hdr_cells2 = table2.rows[0].cells
for i, name in enumerate(headers2):
    hdr_cells2[i].text = name
    set_cell_background(hdr_cells2[i], "1F4E78")
    set_cell_margins(hdr_cells2[i])
    p = hdr_cells2[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

t2_data = [
    ("1. Excel 實機操作", "30%", "5分：10分鐘完勝且抓出所有過磅異常\n3分：完成基本公式但未抓出異常\n1分：無法正確使用 VLOOKUP", "", ""),
    ("2. ERP與過磅邏輯", "25%", "5分：具備嚴謹單據對帳與權限觀念\n3分：懂基本庫存邏輯但防呆意識弱\n1分：隨意變更系統數據", "", ""),
    ("3. 彰濱廠區適應力", "25%", "5分：完全接受廠區環境(過磅/盤點/粉塵)\n3分：可接受但需適應期\n1分：僅接受純辦公室環境", "", ""),
    ("4. 細心度與溝通力", "20%", "5分：對數據高度敏感，溝通得體堅持原則\n3分：溝通良好但偶有粗心\n1分：缺乏耐心或細心度不足", "", ""),
    ("綜合加權總分", "100%", "錄取門檻標準：75 分以上 (滿分 100 分)", "", "")
]

for row_idx, data in enumerate(t2_data, start=1):
    row_cells = table2.rows[row_idx].cells
    for col_idx, text in enumerate(data):
        row_cells[col_idx].text = text
        set_cell_margins(row_cells[col_idx])
        if row_idx == 5:
            set_cell_background(row_cells[col_idx], "D9E1F2")
        elif row_idx % 2 == 1:
            set_cell_background(row_cells[col_idx], "F2F2F2")
        p = row_cells[col_idx].paragraphs[0]
        if col_idx in [1, 3]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for row in table2.rows:
    for idx, width in enumerate(widths2):
        row.cells[idx].width = width

doc.save(r"d:\GOOGLE ANGET\彰濱廠區_助理管理師_全務招募面試與評量手冊.docx")
print("Successfully generated Word document with dual-excel reference!")
