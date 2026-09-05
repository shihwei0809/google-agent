import os
from docx import Document
from docx.shared import Pt

def generate_manual():
    print("開始生成操作手冊...")
    doc = Document()
    
    # 標題
    title = doc.add_heading('AI 教育訓練平台 - 操作手冊', 0)
    title.alignment = 1
    
    # 簡介
    doc.add_heading('1. 系統簡介', level=1)
    doc.add_paragraph('本系統為結合 AI 助理的教育訓練平台。提供學員閱覽教材並隨時與 AI 助教進行問答。')
    
    # 系統安裝
    doc.add_heading('2. 環境安裝與啟動', level=1)
    doc.add_paragraph('1. 點擊執行 setup_env.ps1 進行環境配置。')
    doc.add_paragraph('2. 啟動後端：進入 backend 目錄，執行 uvicorn main:app --reload')
    doc.add_paragraph('3. 啟動前端：進入 frontend 目錄，執行 npm run dev')
    
    # 使用說明
    doc.add_heading('3. 核心功能操作', level=1)
    doc.add_paragraph('【教材區】：左側清單選擇欲閱讀之教材。')
    doc.add_paragraph('【AI 助教】：右側對話框輸入問題，AI 會根據教材上下文給予解答。')
    
    # 儲存
    filename = 'AI教育訓練平台_操作手冊.docx'
    doc.save(filename)
    print(f"手冊已成功生成：{filename}")
    
    # 注意：Windows 環境下轉換 PDF 通常需要依賴 Word 軟體 COM 介面或額外套件，
    # 這裡先產出 docx，PDF 產出可透過手動或安裝 docx2pdf 擴充。

if __name__ == "__main__":
    try:
        generate_manual()
    except ImportError:
        print("缺少 python-docx 套件，請先執行 pip install python-docx")
