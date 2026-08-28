import sys
import re
try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Please install python-docx (pip install python-docx)")
    sys.exit(1)

def convert_md_to_docx(md_path, docx_path):
    document = Document()
    
    style = document.styles['Title']
    font = style.font
    font.name = '敺株?甇??擃?
    font.size = Pt(20)

    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except FileNotFoundError:
        print(f"Error: {md_path} not found.")
        sys.exit(1)

    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('# '):
            p = document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = document.add_heading(line[4:], level=3)
        elif line.startswith('* ') or line.startswith('- '):
            text = line[2:]
            p = document.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)
        elif line.startswith('1. ') or line.startswith('2. '):
            p = document.add_paragraph(style='List Number')
            _add_formatted_text(p, line[3:])
        else:
            p = document.add_paragraph()
            _add_formatted_text(p, line)

    document.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

def _add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sub_part in sub_parts:
                if sub_part.startswith('*') and sub_part.endswith('*'):
                    run = paragraph.add_run(sub_part[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sub_part)

if __name__ == '__main__':
    md_file = '1000m2_Chemical_Pallet_Shuttle_Feasibility_Report_v4.md'
    docx_file = '1000m2_Chemical_Pallet_Shuttle_Feasibility_Report_v4.docx'
    convert_md_to_docx(md_file, docx_file)

