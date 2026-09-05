import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32com.client

def create_manual():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("【系統操作手冊】鴻勝包材管理系統 v13.5 (PWA 獨立 App 版)")
    run_title.font.name = "微軟正黑體"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(13, 110, 253)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("版本：v13.5 (PWA 雙軌版) | 適用設備：Windows 電腦 / Android 手機 / iPhone / 平板")
    run_sub.font.name = "微軟正黑體"
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph("―" * 50)
    
    doc.add_heading("一、 系統簡介與核心三大分頁", level=1)
    doc.add_paragraph("本系統為「鴻勝包材管理系統 v13.5」之 PWA 現代化雙軌版本，具備試算表自動分流 (NSE_Log / Inventory_Log)、自動對位新增欄位與批號發出次數追蹤：")
    
    tabs = [
        ("1. 📦 總表 (Board)", "依類別（鴻勝新桶、客供新桶、特定原料、回收桶、成品、其他包裝耗材）展示即時動態庫存與安全水位警戒。"),
        ("2. 🔍 帳卡 (Query)", "支援歷史帳卡關鍵字即時搜尋、跨月份對帳與依日期自動排序合併明細。"),
        ("3. 📝 作業 (Input)", "進出庫登打、領用單位自動切換、批號剩餘量自動帶出與發出次數即時提示。")
    ]
    for t_name, t_desc in tabs:
        p = doc.add_paragraph()
        r = p.add_run(f"◆ {t_name}：\n")
        r.bold = True
        r.font.color.rgb = RGBColor(13, 110, 253)
        p.add_run(t_desc)

    doc.add_heading("二、 PWA 獨立 App 安裝與啟動指引", level=1)
    steps = [
        ("電腦端 (Chrome/Edge)", "雙擊「啟動PWA本機測試.bat」，點擊網址列右側「安裝應用程式」或頂部「立即安裝」，桌面即產生獨立 App。"),
        ("Android 手機 / PDA", "連線網址後，點擊頂部「立即安裝」或瀏覽器選單「新增至主畫面」。"),
        ("iPhone / iPad (Safari)", "點擊 Safari 底部「分享」圖示 -> 選擇「加入主畫面」即可安裝。")
    ]
    for s_title, s_desc in steps:
        p = doc.add_paragraph()
        r = p.add_run(f"【{s_title}】\n")
        r.bold = True
        p.add_run(s_desc)

    doc_path = r"C:\GOOGLE ANGET\倉庫常用包材管理\2_PWA_App版\鴻勝包材管理系統_操作手冊.docx"
    doc.save(doc_path)
    print(f"Generated Word: {doc_path}")

    pdf_path = r"C:\GOOGLE ANGET\倉庫常用包材管理\2_PWA_App版\鴻勝包材管理系統_操作手冊.pdf"
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc_pdf = word.Documents.Open(os.path.abspath(doc_path))
        doc_pdf.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc_pdf.Close()
        word.Quit()
        print(f"Generated PDF: {pdf_path}")
    except Exception as e:
        print(f"PDF note: {e}")

create_manual()
