import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

DISC_REFERENCE_GUIDE = [
    ("D (Dominance) 支配/主導型", "果斷獨立、目標導向、重效率與結果、具備開創力與領導膽識。", "高階主管、高壓業務、轉型專案PM、開拓型經理人"),
    ("I (Influence) 影響/社交型", "熱情具感染力、善於人際溝通、說服力強、樂觀創造力高。", "第一線 sales、公關發言人、活動企劃、團隊溝通協調者"),
    ("S (Steadiness) 穩健/支援型", "溫和耐心、重視團隊和諧、擅長傾聽、重視穩定與高忠誠度。", "【資材行政】、客戶服務專員、HR 人資後勤、行政專員"),
    ("C (Conscientiousness) 謹慎/分析型", "嚴謹精確、邏輯導向、注重細節品質、客觀講求事實依據。", "【資材現場助理工程師/資材工程師】、數據分析師、財務稽核")
]

DEPT_SPECIFIC_GUIDE = [
    ("資材部 - 現場助理工程師", "CS 型 (C型60% + S型40%)", "高精準度、料號與數據敏銳、耐受現場庫存與備料SOP、零失誤率。"),
    ("資材部 - 資材工程師", "CS 型 (C型70% + S型30%)", "BOM表結構化分析、物料供需規劃、嚴謹邏輯與數據管控。"),
    ("資材部 - 資材行政專員", "SC 型 (S型60% + C型40%)", "跨部門溝通協調、耐心後勤支援、表單與 ERP 輸入精準度。")
]

def generate_interview_report_docx(record: dict, output_path: str):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    report = record.get("report") or {}
    target_dept = record.get("target_dept", "資材部 (現場助理工程師/工程師/行政)")

    # 主標題
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("AI 面試語音特質與職務適性評估報告")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(30, 41, 59)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub_p.add_run(f"應徵目標部門：{target_dept}  |  評估時間：{record.get('created_at', 'N/A')}")
    run_sub.font.name = 'Microsoft JhengHei'
    run_sub.font.size = Pt(10)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(79, 70, 229)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 1. 概要表格
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    cell_00 = table.cell(0, 0)
    cell_00.text = "DISC 本次判定人格"
    set_cell_background(cell_00, "F1F5F9")

    cell_01 = table.cell(0, 1)
    cell_01.text = report.get("disc_type", "N/A")
    set_cell_background(cell_01, "EEF2FF")

    cell_10 = table.cell(1, 0)
    cell_10.text = "對話核心摘要"
    set_cell_background(cell_10, "F1F5F9")

    cell_11 = table.cell(1, 1)
    cell_11.text = report.get("candidate_summary", "N/A")

    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Microsoft JhengHei'
                    run.font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # === 資材與特定部門 DISC 指南 ===
    h_dept = doc.add_heading("一、 資材與特定部門理想 DISC 人格與適性對照表", level=2)
    h_dept.runs[0].font.name = 'Microsoft JhengHei'
    h_dept.runs[0].font.color.rgb = RGBColor(79, 70, 229)

    dept_table = doc.add_table(rows=4, cols=3)
    dept_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    c0 = dept_table.cell(0, 0); c0.text = "目標職務"
    c1 = dept_table.cell(0, 1); c1.text = "黃金 DISC 類型"
    c2 = dept_table.cell(0, 2); c2.text = "關鍵能力與適性要求"
    for c in (c0, c1, c2):
        set_cell_background(c, "312E81")
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        c.paragraphs[0].runs[0].font.bold = True

    for idx, (title, disc_req, desc) in enumerate(DEPT_SPECIFIC_GUIDE, start=1):
        c0 = dept_table.cell(idx, 0); c0.text = title
        c1 = dept_table.cell(idx, 1); c1.text = disc_req
        c2 = dept_table.cell(idx, 2); c2.text = desc
        if idx % 2 == 1:
            set_cell_background(c0, "F8FAFC")
            set_cell_background(c1, "F8FAFC")
            set_cell_background(c2, "F8FAFC")
        for c in (c0, c1, c2):
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Microsoft JhengHei'; r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # === DISC 全類型 HR 參考說明表 ===
    h_disc = doc.add_heading("二、 DISC 四大核心人格類型通用對照表", level=2)
    h_disc.runs[0].font.name = 'Microsoft JhengHei'
    h_disc.runs[0].font.color.rgb = RGBColor(79, 70, 229)

    disc_table = doc.add_table(rows=5, cols=3)
    disc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers_disc = [("DISC 人格類型", "行為風格與典型特徵", "典型適合職務類型")]
    for col_idx, h_text in enumerate(headers_disc[0]):
        c = disc_table.cell(0, col_idx)
        c.text = h_text
        set_cell_background(c, "4F46E5")
        p = c.paragraphs[0]
        p.runs[0].font.name = 'Microsoft JhengHei'
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.bold = True

    for idx, (disc_name, desc, jobs) in enumerate(DISC_REFERENCE_GUIDE, start=1):
        c0 = disc_table.cell(idx, 0)
        c1 = disc_table.cell(idx, 1)
        c2 = disc_table.cell(idx, 2)
        c0.text = disc_name
        c1.text = desc
        c2.text = jobs
        
        current_disc = report.get("disc_type", "")
        if disc_name[0] in current_disc:
            set_cell_background(c0, "FEF3C7")
            set_cell_background(c1, "FEF3C7")
            set_cell_background(c2, "FEF3C7")
        elif idx % 2 == 1:
            set_cell_background(c0, "F8FAFC")
            set_cell_background(c1, "F8FAFC")
            set_cell_background(c2, "F8FAFC")

        for c in (c0, c1, c2):
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = 'Microsoft JhengHei'
                    r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 3. 聲學與溝通表達觀察
    h2 = doc.add_heading("三、 語音聲理與溝通風格觀察", level=2)
    h2.runs[0].font.name = 'Microsoft JhengHei'
    h2.runs[0].font.color.rgb = RGBColor(79, 70, 229)

    p_comm = doc.add_paragraph()
    r = p_comm.add_run("• 溝通表達風格：")
    r.bold = True
    p_comm.add_run(report.get("communication_style", "N/A"))

    p_acous = doc.add_paragraph()
    r = p_acous.add_run("• 聲學與語速細節：")
    r.bold = True
    p_acous.add_run(report.get("acoustic_observation", "N/A"))

    # 4. Big Five 五大人格量化指標
    h3 = doc.add_heading("四、 Big Five 五大人格量化評估 (OCEAN)", level=2)
    h3.runs[0].font.name = 'Microsoft JhengHei'
    h3.runs[0].font.color.rgb = RGBColor(79, 70, 229)

    bf = report.get("big_five") or {}
    bf_table = doc.add_table(rows=6, cols=2)
    bf_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [("人格維度 (Trait)", "評分 (1-100)"), 
               ("經驗開放性 (Openness)", f"{bf.get('openness', 0)}%"),
               ("盡責性 (Conscientiousness)", f"{bf.get('conscientiousness', 0)}%"),
               ("外向性 (Extraversion)", f"{bf.get('extraversion', 0)}%"),
               ("宜人性 (Agreeableness)", f"{bf.get('agreeableness', 0)}%"),
               ("情緒穩定度 (Emotional Stability)", f"{bf.get('emotional_stability', 0)}%")]

    for idx, (label, val) in enumerate(headers):
        c0 = bf_table.cell(idx, 0)
        c1 = bf_table.cell(idx, 1)
        c0.text = label
        c1.text = val
        if idx == 0:
            set_cell_background(c0, "4F46E5")
            set_cell_background(c1, "4F46E5")
            for p in c0.paragraphs: p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            for p in c1.paragraphs: p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        elif idx % 2 == 1:
            set_cell_background(c0, "F8FAFC")
            set_cell_background(c1, "F8FAFC")

        for c in (c0, c1):
            set_cell_margins(c, top=80, bottom=80, left=150, right=150)
            for p in c.paragraphs:
                for run in p.runs: run.font.name = 'Microsoft JhengHei'

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 5. 推薦與不推薦職能
    h4 = doc.add_heading("五、 職能適性契合度分析", level=2)
    h4.runs[0].font.name = 'Microsoft JhengHei'
    h4.runs[0].font.color.rgb = RGBColor(79, 70, 229)

    doc.add_paragraph(f"【推薦契合職務 (針對 {target_dept})】").runs[0].bold = True
    for j in report.get("recommended_jobs") or []:
        p = doc.add_paragraph()
        r = p.add_run(f"✓ {j.get('title')} ({j.get('score')}% 契合)：")
        r.bold = True
        r.font.color.rgb = RGBColor(16, 185, 129)
        p.add_run(j.get("reason", ""))

    doc.add_paragraph("【較不適合職務 (Low Fit)】").runs[0].bold = True
    for j in report.get("unsuitable_jobs") or []:
        p = doc.add_paragraph()
        r = p.add_run(f"✕ {j.get('title')} ({j.get('score')}% 契合)：")
        r.bold = True
        r.font.color.rgb = RGBColor(244, 63, 94)
        p.add_run(j.get("reason", ""))

    # 6. 主管帶領建議與二面追問
    h5 = doc.add_heading("六、 主管帶領建議與二面追問題目", level=2)
    h5.runs[0].font.name = 'Microsoft JhengHei'
    h5.runs[0].font.color.rgb = RGBColor(79, 70, 229)

    p_mgmt = doc.add_paragraph()
    p_mgmt.add_run("【團隊融入與帶領建議】\n").bold = True
    p_mgmt.add_run(report.get("management_advice", "N/A"))

    p_q = doc.add_paragraph()
    p_q.add_run("【建議二面追問題目】").bold = True
    for q in report.get("followup_questions") or []:
        doc.add_paragraph(f"• {q}")

    # 7. 附錄：完整對話逐字稿
    h6 = doc.add_heading("七、 附錄：面試語音完整逐字稿", level=2)
    h6.runs[0].font.name = 'Microsoft JhengHei'
    h6.runs[0].font.color.rgb = RGBColor(79, 70, 229)

    p_trans = doc.add_paragraph()
    run_t = p_trans.add_run(record.get("transcript", "(無逐字稿)"))
    run_t.font.name = 'Microsoft JhengHei'
    run_t.font.size = Pt(9.5)
    run_t.font.color.rgb = RGBColor(71, 85, 105)

    doc.save(output_path)
    return output_path
