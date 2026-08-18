import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import win32com.client

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_interview_manual_docx(docx_path):
    doc = docx.Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 標題
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("🎙️ AI 面試語音特質與資材適性分析系統\n人員操作手冊 (User Manual)")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
    p_title.paragraph_format.space_after = Pt(15)

    # 副標題
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_sub = p_sub.add_run("適用對象：HR 人資團隊、資材與各部門主管 | 版本：v2.0")
    r_sub.font.size = Pt(9.5)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p_sub.paragraph_format.space_after = Pt(20)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Microsoft JhengHei'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Microsoft JhengHei'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            r_b.font.bold = True
            r_b.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
        p.add_run(text)
        return p

    def add_tip(text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "E8F5E9")
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("💡 " + text)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 1. 系統簡介與核心定位
    add_h1("一、 系統簡介與核心定位")
    p1 = doc.add_paragraph("本系統專為人資 (HR) 團隊與部門主管設計，結合 Gemini 語言 AI 模型與多重人格分析架構 (DISC 與 Big Five 五大人格)。提供「事前履歷精準評估」與「現場錄音逐字稿轉寫評估」，協助主管快速抓出應徵者優勢、技能落差與適性型態。")
    p1.paragraph_format.space_after = Pt(6)
    add_bullet("支援上傳 PDF、圖片 (104 履歷截圖/掃描) 或文字，比對 27 個實體職缺清單，產出匹配分數、技能落差與 4~6 個結構化提問指南。", "1. 事前履歷 AI 分析 (Pre-interview)：")
    add_bullet("每 5 分鐘背景自動備份拼接，1 小時長面試零等待落盤；自動產出逐字稿、DISC 與 Big Five 適性評估。", "2. 現場錄音分段備份轉寫：")
    add_bullet("現場助理工程師 (黃金 CS 型)、資材工程師 (黃金 CS 型)、資材行政專員 (黃金 SC 型)。", "3. 資材專屬黃金適性指標：")
    add_bullet("將事前評估與面試結果一鍵產出為符合企業規範之 Word (.docx) 與 Excel (.xlsx) 報告。", "4. 雙格式報告導出：")

    # 2. 一鍵啟動與登入
    add_h1("二、 一鍵啟動與區網共用")
    add_h2("1. 快速啟動方式")
    doc.add_paragraph("在專案資料夾中，直接雙擊執行：")
    p_cmd = doc.add_paragraph()
    p_cmd.paragraph_format.left_indent = Inches(0.3)
    r_cmd = p_cmd.add_run("👉 雙擊點我啟動面試語音AI分析系統.bat")
    r_cmd.font.bold = True
    r_cmd.font.size = Pt(12)
    r_cmd.font.color.rgb = RGBColor(0x21, 0x96, 0xF3)
    p_cmd.paragraph_format.space_after = Pt(6)
    add_tip("系統會自動檢查環境套件，並自動為您打開預設瀏覽器連至網頁介面。")

    add_h2("2. 區域網路 (LAN) 同事共用連線")
    doc.add_paragraph("系統啟動時會在主控台黑視窗中顯示本機實體 IP 位址：")
    add_bullet("http://localhost:8000", "• 本機開啟：")
    add_bullet("http://[您的本機IP]:8000 (例如 http://192.168.1.100:8000)，同辦公室同仁無須安裝即可直接登入存取！", "• 區網同仁開啟：")

    # 3. 核心模組操作說明
    add_h1("三、 核心模組操作說明")
    
    add_h2("模組 1：事前履歷 AI 分析與提問指南 (Pre-interview)")
    add_bullet("點擊網頁頂部「事前履歷分析」分頁。", "步驟 1：")
    add_bullet("拖曳上傳 104 履歷 PDF 檔、圖片截圖/掃描檔，或直接將履歷文字貼入文字方塊中。", "步驟 2：")
    add_bullet("在下拉選單選擇應徵職缺（包含【彰濱廠區】及【高雄廠】共 27 個實體刊登職務與 AI 範本職缺）。", "步驟 3：")
    add_bullet("點擊「🚀 開始 AI 履歷剖析」，系統將在一鍵分析後提供：學經歷匹配度分項（0-100分）、三大核心優勢、技能落差風險、通勤轉折評估，以及 4~6 個面試官觀察重點指南。", "步驟 4：")
    add_bullet("點擊「📥 下載 Word 事前評估報告」，即可產出標準 .docx 檔案。", "步驟 5：")

    add_h2("模組 2：現場面試錄音與 DISC 適性評估 (Live Interview)")
    add_bullet("點擊「現場錄音分析」分頁。", "步驟 1：")
    add_bullet("選擇應徵部門適性範本（如『資材工程師 - 黃金 CS 型』）。", "步驟 2：")
    add_bullet("點擊「🎙️ 開始錄音」，系統會自動每 5 分鐘背景分段備份（防止長時間錄音因瀏覽器崩潰或斷電丟失資料）。", "步驟 3：")
    add_bullet("面試結束點擊「⏹️ 結束錄音並分析」，AI 自動進行語音轉寫逐字稿、情緒標註、DISC 特質評估與 Big Five 五大人格指標比對。", "步驟 4：")
    add_bullet("點擊「📥 導出 Word 面試報告」或「📊 導出 Excel 面試報告」匯出完整檔案。", "步驟 5：")

    add_h2("模組 3：履歷圖片 / PDF 直接轉換為 Excel 試算表")
    doc.add_paragraph("若您只希望將大量的 PDF/圖片履歷直接萃取轉換為 Excel 試算表：")
    add_bullet("點擊「履歷轉 Excel」按鈕，上傳 PDF 或圖片檔，系統會將聯絡資料、學經歷、歷練職掌與自傳 100% 完整轉寫為標準 `.xlsx` 試算表檔。")

    # 4. 歷史紀錄管理與斷網重辨識
    add_h1("四、 歷史紀錄管理與斷網補救")
    add_bullet("點擊「歷史紀錄」標籤，可查詢過去分析過的所有履歷評估與現場錄音檔案。", "1. 歷史查詢：")
    add_bullet("對於已完成的面試紀錄，隨時可重新點擊下載 Word 或 Excel 報告，或一鍵刪除舊紀錄。", "2. 補下載/刪除：")
    add_bullet("若錄音過程中突然斷網，音檔已完整備份於 data/audios/ 資料夾中，點擊「離線音檔重辦識」即可由系統補重新進行 AI 解析。", "3. 斷網補救重辨識：")

    # 5. 多金鑰輪替與模型備援
    add_h1("五、 多金鑰輪替與 AI 模型備援設定")
    doc.add_paragraph("本系統預設內建 Gemini 多 API Key 白嫖與瀑布流降階備援機制：")
    add_bullet("可在網頁右上角「⚙️ API Key 設定」中貼入多組 Gemini API 金鑰（以換行分隔），系統會自動進行 Round-Robin 輪替，避免達到限額 (429 Rate Limit)。", "• 多金鑰輪替 (Multi-Key)：")
    add_bullet("當主要模型繁忙時，系統會自動按順序嘗試：gemini-3.6-flash ➔ gemini-3.5-flash ➔ gemini-3.1-flash-lite ➔ gemini-2.5-flash，確保服務永遠穩定不中斷！", "• 模型瀑布流 (Cascade Fallback)：")

    # 6. FAQ
    add_h1("六、 常見問題 FAQ")
    add_bullet("請確認電腦是否已連接麥克風，並在瀏覽器跳出提示時選擇『允許存取麥克風』。", "Q1：錄音按鈕點擊無反應？ ➔ ")
    add_bullet("長錄音檔採用背景 5 分鐘分段備份，錄音過程中瀏覽器關閉亦會保存已錄製之音檔，可在『歷史紀錄』中恢復離線辨識。", "Q2：錄音中途關閉網頁怎麼辦？ ➔ ")
    add_bullet("請至網頁右上角『⚙️ API Key 設定』確認金鑰是否過期或限額，建議貼入 2~3 組免費金鑰進行輪替。", "Q3：跳出 API Rate Limit 限額警告？ ➔ ")

    doc.save(docx_path)
    print(f"Word generated at: {docx_path}")

def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        word.Quit()
        print(f"PDF generated at: {pdf_path}")
    except Exception as e:
        print(f"Failed to convert PDF: {e}")

if __name__ == "__main__":
    base_dir = r"d:\GOOGLE ANGET\第三類_AI代理與指南企劃\interview_analyzer"
    docx_file = os.path.join(base_dir, "AI面試語音特質與資材適性分析系統_操作手冊.docx")
    pdf_file = os.path.join(base_dir, "AI面試語音特質與資材適性分析系統_操作手冊.pdf")
    
    create_interview_manual_docx(docx_file)
    convert_docx_to_pdf(docx_file, pdf_file)
