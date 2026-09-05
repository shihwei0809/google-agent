import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_manual():
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("【系統操作手冊】N系列出貨條碼核對系統 (PWA 獨立 App 版)")
    run_title.font.name = "微軟正黑體"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(25, 118, 210)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("版本號：v1.0 (PWA App 版) | 適用設備：Windows 電腦 / Android 手機 / iPhone / 手持 PDA")
    run_sub.font.name = "微軟正黑體"
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph("―" * 50)
    
    # Section 1
    h1 = doc.add_heading("一、 系統簡介與 PWA 技術特色", level=1)
    doc.add_paragraph("本系統為「N系列槽車出貨條碼核對」之現代化 PWA（Progressive Web App）應用程式版本。具備以下三大核心優勢：")
    
    p = doc.add_paragraph()
    p.add_run("1. 免上架商店，一鍵安裝：").bold = True
    p.add_run(" 無需透過 Google Play 或 App Store，透過瀏覽器即可一鍵將系統安裝至桌面或手機主畫面。\n")
    p.add_run("2. 獨立全螢幕運行：").bold = True
    p.add_run(" 安裝後點擊專屬圖示開啟，無瀏覽器網址列與多餘按鈕，提供 100% 原生 App 沉浸體驗。\n")
    p.add_run("3. 離線快取與相機調用：").bold = True
    p.add_run(" 內建 Service Worker 背景快取機制，即使現場網路訊號微弱亦能快速啟動，並原生支援手機相機條碼掃描。")
    
    # Section 2
    h2 = doc.add_heading("二、 電腦端 (Windows / Mac) 安裝與操作步驟", level=1)
    
    steps_pc = [
        ("步驟 1：啟動本機測試伺服器", "雙擊資料夾內的「啟動PWA本機測試.bat」，黑視窗將自動偵測可用 Port 與本機 IP 並開啟瀏覽器。"),
        ("步驟 2：點擊瀏覽器安裝圖示", "使用 Google Chrome 或 Microsoft Edge 開啟網址，網址列最右端會出現「安裝應用程式 (螢幕圖示)」或點擊頁面上方藍色「立即安裝」按鈕。"),
        ("步驟 3：確認安裝至桌面", "在彈出的確認視窗點擊「安裝」，電腦桌面將自動建立「N系列出貨核對」獨立捷徑圖示。"),
        ("步驟 4：日常獨立啟動", "往後只需雙擊桌面圖示，系統即以乾淨的獨立原生視窗開啟，無需再手動輸入網址。")
    ]
    for title, desc in steps_pc:
        p = doc.add_paragraph()
        r = p.add_run(f"【{title}】\n")
        r.bold = True
        r.font.color.rgb = RGBColor(25, 118, 210)
        p.add_run(desc)
        
    # Section 3
    h3 = doc.add_heading("三、 手機與手持 PDA (Android / iPhone) 安裝步驟", level=1)
    
    steps_mobile = [
        ("Android 手機 / 手持條碼 PDA (Chrome 瀏覽器)", "1. 開啟 Chrome 連線至伺服器網址。\n2. 點擊頁面頂部「立即安裝」橫幅，或點擊右上角「...」選單 -> 選擇「安裝應用程式」或「加到主畫面」。\n3. 手機桌面即出現 App 圖示，點開即可全螢幕掃描條碼。"),
        ("iPhone / iPad (Safari 瀏覽器)", "1. 使用 Safari 開啟系統網址。\n2. 點擊底部工具列正中央的「分享 (方框向上箭頭)」按鈕。\n3. 往下滑動選擇「加入主畫面 (Add to Home Screen)」並點擊右上角「新增」。\n4. 桌面即產生專屬 App 圖示。")
    ]
    for title, desc in steps_mobile:
        p = doc.add_paragraph()
        r = p.add_run(f"◆ {title}\n")
        r.bold = True
        r.font.color.rgb = RGBColor(0, 150, 136)
        p.add_run(desc)

    # Section 4 FAQ
    doc.add_heading("四、 常見問題排除 (FAQ)", level=1)
    faqs = [
        ("Q1：為什麼 Chrome 網址列沒有出現安裝按鈕？", "A1：請確認是否已點擊過安裝（若已安裝則不會重複顯示），或點擊頁面頂部的「立即安裝」藍色按鈕即可觸發。"),
        ("Q2：手機相機無法掃描條碼？", "A2：首次開啟時請務必在瀏覽器彈出的權限視窗點擊「允許使用相機」。若誤按拒絕，請至手機設定 -> 應用程式 -> 瀏覽器 -> 權限中重新開啟相機。"),
        ("Q3：更換電腦時如何遷移？", "A3：PWA 專案所有設定均包含在資料夾內，直接複製整個資料夾至新電腦並執行「啟動PWA本機測試.bat」即可立即使用。")
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        r_q = p.add_run(f"{q}\n")
        r_q.bold = True
        p.add_run(a)

    doc_path = r"C:\GOOGLE ANGET\第二類_生產管理與API串接\N系列BARCODE出貨核對_PWA\N系列BARCODE出貨核對_操作手冊.docx"
    doc.save(doc_path)
    print(f"Generated Word Manual: {doc_path}")

    # Convert to PDF
    pdf_path = r"C:\GOOGLE ANGET\第二類_生產管理與API串接\N系列BARCODE出貨核對_PWA\N系列BARCODE出貨核對_操作手冊.pdf"
    try:
        from docx2pdf import convert
        convert(doc_path, pdf_path)
        print(f"Generated PDF Manual: {pdf_path}")
    except Exception as e:
        print(f"PDF convert note: {e}")

create_manual()
