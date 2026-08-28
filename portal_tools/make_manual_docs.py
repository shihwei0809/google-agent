import os
import sys
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import win32com.client
import pythoncom
import shutil

try:
    sys.stdout.reconfigure(encoding='utf-8')
except Exception:
    pass

def remove_emojis(text):
    """Remove all emoji characters to prevent empty boxes (□) in Word."""
    emoji_pattern = re.compile(
        r'[\U00010000-\U0010ffff]|[\u2600-\u27bf]|[\u2300-\u23ff]|[\u2b50-\u2b55]|[\ufe00-\ufe0f]',
        flags=re.UNICODE
    )
    return emoji_pattern.sub('', text)

def set_cell_background(cell, fill_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def parse_markdown_table(lines):
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|") or not stripped.endswith("|"):
            continue
        cells = [c.strip() for c in stripped.split("|")[1:-1]]
        if all(re.match(r'^:?-+:?$', c) for c in cells):
            continue
        rows.append(cells)
    return rows

def create_styled_docx(md_path, docx_path):
    md_dir = os.path.dirname(os.path.abspath(md_path))
    with open(md_path, 'r', encoding='utf-8') as f:
        raw_text = f.read()

    md_text = remove_emojis(raw_text)

    doc = Document()
    
    # Page Margins: 0.8 inch
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft JhengHei'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    lines = md_text.splitlines()
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        parsed_rows = parse_markdown_table(table_lines)
        if parsed_rows:
            num_cols = max(len(r) for r in parsed_rows)
            table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            
            for row_idx, row_data in enumerate(parsed_rows):
                is_header = (row_idx == 0)
                for col_idx in range(num_cols):
                    cell = table.cell(row_idx, col_idx)
                    val = row_data[col_idx] if col_idx < len(row_data) else ""
                    
                    if is_header:
                        set_cell_background(cell, "1A237E") # Navy Blue
                        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        run = p.add_run(val)
                        run.font.bold = True
                        run.font.name = 'Microsoft JhengHei'
                        run.font.size = Pt(9.5)
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    else:
                        bg_hex = "F8F9FA" if row_idx % 2 == 1 else "FFFFFF"
                        set_cell_background(cell, bg_hex)
                        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        run = p.add_run(val)
                        run.font.name = 'Microsoft JhengHei'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        table_lines = []

    for line in lines:
        stripped = line.strip()

        # Handle Code Block
        if stripped.startswith("```"):
            if in_code_block:
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = table.cell(0, 0)
                set_cell_background(cell, "F8F9FA")
                set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                run = p.add_run("\n".join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
                
                doc.add_paragraph().paragraph_format.space_after = Pt(4)
                code_lines = []
                in_code_block = False
            else:
                if in_table:
                    flush_table()
                    in_table = False
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Handle Table Lines
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            table_lines.append(stripped)
            continue
        elif in_table:
            flush_table()
            in_table = False

        if not stripped:
            continue

        # Horizontal Rule
        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run("―" * 45)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            continue

        # Image Tag: ![caption](rel_path)
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', stripped)
        if img_match:
            caption = img_match.group(1).strip()
            raw_img_rel = img_match.group(2).strip()
            
            # Resolve image path
            abs_img_path = os.path.normpath(os.path.join(md_dir, raw_img_rel))
            if os.path.exists(abs_img_path):
                # Add image
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(8)
                p_img.paragraph_format.space_after = Pt(2)
                run_img = p_img.add_run()
                run_img.add_picture(abs_img_path, width=Inches(5.8))
                
                # Add caption
                if caption:
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_before = Pt(0)
                    p_cap.paragraph_format.space_after = Pt(8)
                    run_cap = p_cap.add_run(f"圖：{caption}")
                    run_cap.font.name = 'Microsoft JhengHei'
                    run_cap.font.size = Pt(8.5)
                    run_cap.font.italic = True
                    run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            else:
                print(f"[WARN] Image not found: {abs_img_path}")
            continue

        # Heading 1
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(stripped[2:].strip())
            run.font.name = 'Microsoft JhengHei'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
            continue

        # Heading 2
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[3:].strip())
            run.font.name = 'Microsoft JhengHei'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
            continue

        # Heading 3
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[4:].strip())
            run.font.name = 'Microsoft JhengHei'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x69, 0x5C)
            continue

        # Blockquote (starts with >)
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            text = stripped[1:].strip()
            run = p.add_run(text)
            run.font.name = 'Microsoft JhengHei'
            run.font.size = Pt(9.5)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # Bullet lists (- or *)
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            parse_styled_inline(p, stripped[2:].strip())
            continue

        # Numbered lists (1. 2. 3.)
        if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in ('.', '、'):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            parse_styled_inline(p, stripped[2:].strip())
            continue

        # Regular Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        parse_styled_inline(p, stripped)

    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f"[OK] DOCX generated: {docx_path}")

def parse_styled_inline(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|\`.*?\`|\[.*?\]\(.*?\))', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Microsoft JhengHei'
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
        elif part.startswith("[") and "](" in part:
            m = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.font.name = 'Microsoft JhengHei'
                run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
                run.underline = True
            else:
                run = paragraph.add_run(part)
                run.font.name = 'Microsoft JhengHei'
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Microsoft JhengHei'

def convert_docx_to_pdf(docx_path, pdf_path):
    abs_docx = os.path.abspath(docx_path)
    abs_pdf = os.path.abspath(pdf_path)
    
    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    
    try:
        doc = word.Documents.Open(abs_docx)
        doc.SaveAs(abs_pdf, FileFormat=17) # 17 = wdExportFormatPDF
        doc.Close()
        print(f"[OK] PDF generated: {abs_pdf}")
    except Exception as e:
        print(f"[ERROR] PDF conversion: {e}")
    finally:
        word.Quit()
        pythoncom.CoUninitialize()

if __name__ == "__main__":
    proj_dir = r"C:\GOOGLE ANGET\三合一單網頁架機伺服器"
    md_file = os.path.join(proj_dir, "操作說明書.md")
    docx_file = os.path.join(proj_dir, "三合一單網頁架機伺服器_操作說明書.docx")
    pdf_file = os.path.join(proj_dir, "三合一單網頁架機伺服器_操作說明書.pdf")
    
    docx_copy = os.path.join(proj_dir, "操作說明書.docx")
    pdf_copy = os.path.join(proj_dir, "操作說明書.pdf")

    create_styled_docx(md_file, docx_file)
    shutil.copy2(docx_file, docx_copy)
    
    convert_docx_to_pdf(docx_file, pdf_file)
    if os.path.exists(pdf_file):
        shutil.copy2(pdf_file, pdf_copy)
