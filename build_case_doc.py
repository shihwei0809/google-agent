import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import win32com.client

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def generate_harassment_doc(docx_path, img_factory, img_office):
    doc = docx.Document()
    
    # 邊界設定
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # 設定預設字型
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 主標題
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("【職場霸凌與性騷擾案例分析與法規辨識手冊】")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
    p_title.paragraph_format.space_after = Pt(6)

    # 副標題
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("廠區作業與辦公室情境實務案例彙整評析 (男同儕與職場性別友善專題)")
    r_sub.font.size = Pt(11)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p_sub.paragraph_format.space_after = Pt(20)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Microsoft JhengHei'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Microsoft JhengHei'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            r_b.font.bold = True
            r_b.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
        p.add_run(text)
        return p

    def add_card_box(title, text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "FFF8E7")
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r_t = p.add_run("◆ " + title + "\n")
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        p.add_run(text)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_image_centered(img_path, caption=""):
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run()
            run.add_picture(img_path, width=Inches(6.2))
            
            if caption:
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(10)
                r_cap = p_cap.add_run(caption)
                r_cap.font.size = Pt(9.5)
                r_cap.font.italic = True
                r_cap.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 1. 前言與概念介紹
    add_h1("一、 前言與概念簡介")
    p_intro = doc.add_paragraph("在職場環境中，不當行為往往兼具「職場霸凌」與「性騷擾」的雙重成分。傳統觀念常誤以為性騷擾僅發生於異性之間或女性受害者，然而在實務上，同性同儕之間、假借玩笑名義、或是針對「性別氣質」的貶抑，同樣構成違反《性別平等工作法》與《勞動基準法》的違法行為。本手冊將兩大真實情境（廠區作業現場與辦公室社交）共 6 大案例進行綜合對比分析。")
    p_intro.paragraph_format.space_after = Pt(10)

    # 2. 第一部份：廠區作業現場情境案例
    add_h1("二、 第一部分：廠區作業現場情境案例評析")
    doc.add_paragraph("本部分聚焦於更衣室、作業區、交接班會議、吸菸區等廠區現場常見之肢體與言語侵犯行為。")
    
    # 插入圖一：廠區作業人員案例示意圖
    add_image_centered(img_factory, "圖 1：男性間現場工作人員：職場霸凌與性騷擾案例示意圖")

    # 廠區表格
    table1 = doc.add_table(rows=1, cols=6)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = False
    
    hdr1 = table1.rows[0].cells
    hdr1_titles = ["案例類型", "發生場景", "行為樣態", "霸凌成分", "性騷擾成分", "加害者常見藉口"]
    hdr1_widths = [Inches(1.1), Inches(1.0), Inches(1.6), Inches(1.3), Inches(1.3), Inches(1.1)]
    
    for i, title in enumerate(hdr1_titles):
        cell = hdr1[i]
        cell.width = hdr1_widths[i]
        set_cell_background(cell, "002060")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9.5)

    data1 = [
        (
            "【案例一】\n假借「玩笑」的肢體侵犯",
            "更衣室、狹窄作業區（如高架或槽體內）",
            "資深員工以「迎新、好玩」為由，對新進男員工進行「阿魯巴」、強行拉扯衣物，或是故意在交錯走動時，蓄意碰觸臀部與下體。",
            "利用資歷與人數優勢（多對一）使新人不敢反抗，形成權力霸凌。",
            "違反意願的私密部位觸碰，構成「肢體性騷擾」。",
            "「大家都是男人，摸一下會少一塊肉嗎？」"
        ),
        (
            "【案例二】\n針對「性別氣質」的言語羞辱",
            "交接班會議、搬運現場",
            "領班針對體能較弱或性格溫和的男員工，公開辱罵「娘砲」、「沒卵葩」等貶抑字眼，並刻意分配極度不合理的粗重工作。",
            "聯合其他同事孤立受害者，並濫用職權給予不當勞動條件，屬於排擠與針對性霸凌。",
            "針對性別特徵與氣質進行貶抑，違反《性別平等工作法》，構成「性別氣質騷擾」。",
            "「做粗活還這麼嬌弱，我是為了鍛鍊他。」"
        ),
        (
            "【案例三】\n權力展現的「敵意環境」建立",
            "吸菸區、現場人員通訊群組",
            "小主管常在休息時間強迫展示色情影片，並當眾詢問員工性生活細節。若員工反感，便會被嘲笑，甚至在排班或考績上遭到刁難。",
            "利用排班與考評權力進行威脅與報復，屬於典型的職權霸凌。",
            "散布色情圖文並製造讓人感到冒犯、恐懼的工作環境，構成「敵意環境性騷擾」。",
            "「開個黃腔而已，也太開不起玩笑了。」"
        )
    ]

    for r_idx, row in enumerate(data1):
        row_cells = table1.add_row().cells
        bg_col = "F9F9F9" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = row_cells[c_idx]
            cell.width = hdr1_widths[c_idx]
            set_cell_background(cell, bg_col)
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(9)
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
            elif c_idx == 5:
                r.font.italic = True
                r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 3. 第二部份：辦公室與社交情境案例
    add_h1("三、 第二部分：辦公室與社交情境案例評析")
    doc.add_paragraph("本部分聚焦於辦公室茶水間、通訊軟體群組、休息區及應酬聚會等軟性社交情境中的越界行為。")

    # 插入圖二：辦公室 Bro Culture 案例示意圖
    add_image_centered(img_office, "圖 2：同事間（男性）職場霸凌與性騷擾案例示範 - 辦公室 Bro Culture 隱蔽騷擾與霸凌")

    table2 = doc.add_table(rows=1, cols=6)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.autofit = False
    
    hdr2 = table2.rows[0].cells
    for i, title in enumerate(hdr1_titles):
        cell = hdr2[i]
        cell.width = hdr1_widths[i]
        set_cell_background(cell, "1B5E20")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9.5)

    data2 = [
        (
            "【案例一】\n通訊軟體與言語的「黃色玩笑」",
            "辦公室茶水間、部門 LINE 群組",
            "同事常在群組傳送色情梗圖，或在茶水間公開拿某位男同事的「性生活」、「單身狀態」或「性別氣質（如：太斯文）」開黃腔嘲笑。",
            "透過公開嘲弄來建立自己在辦公室的社交優勢地位，屬於言語霸凌。",
            "針對他人的性特徵、性別氣質進行嘲笑，並散布令人感到冒犯的圖文，構成「言語與視覺性騷擾」。",
            "「這只是男生之間的幹話，幹嘛這麼嚴肅？」"
        ),
        (
            "【案例二】\n假借放鬆的「肢體越界」",
            "辦公座位區、員工休息區",
            "以「看你最近壓力很大」為由，未經同意強行幫男同事「按摩」肩頸，或在走道交錯時故意拍打對方臀部、大腿作為打招呼的方式。",
            "無視他人身體界線，利用體型或性格優勢強迫對方接受肢體接觸，屬於行為霸凌。",
            "違反當事人意願，觸碰具有性意涵或讓人感到不適的身體部位，構成「肢體性騷擾」。",
            "「大家都是兄弟，抓一下肩膀又不會懷孕。」"
        ),
        (
            "【案例三】\n「兄弟幫」的社交排擠與針對",
            "跨部門會議、下班後的應酬聚會",
            "辦公室內形成特定「兄弟小團體」。若某位男同事不參與他們的粗俗玩笑，就會在專案上被刻意隱瞞資訊，或在會議上被集體貶低專業意見。",
            "刻意孤立特定對象，並阻礙其正常執行工作任務，屬於典型的關係霸凌與職場排擠。",
            "營造出「不加入開黃腔就不合群」的工作氛圍，構成「敵意環境性騷擾」。",
            "「他自己不合群、太難相處，我們才不想跟他合作。」"
        )
    ]

    for r_idx, row in enumerate(data2):
        row_cells = table2.add_row().cells
        bg_col = "F9F9F9" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = row_cells[c_idx]
            cell.width = hdr1_widths[c_idx]
            set_cell_background(cell, bg_col)
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(9)
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
            elif c_idx == 5:
                r.font.italic = True
                r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 4. 法規剖析與企業預防機制
    add_h1("四、 關鍵法規辨識與企業處置建議")
    add_h2("1. 法規三大判別要項")
    doc.add_paragraph("根據《性別平等工作法》與《職業安全衛生法》：")
    add_bullet("性騷擾的成立「不限於異性之間」，且「不以加害者是否有意圖為限」，只要違反當事人主觀意願並感到冒犯，即構成性騷擾。", "• 意願優先原則：")
    add_bullet("針對「娘砲」、「太斯文」等性別刻板印象羞辱，屬於性別氣質騷擾，企業不得默許。", "• 性別氣質保護：")
    add_bullet("主管或同儕散布色情梗圖、建立不合群即排擠的「兄弟幫」文化，構成敵意環境。", "• 敵意環境禁止：")

    add_h2("2. 企業防治處置四步驟 ( Statutory 4-Step Standard Protocol )")
    doc.add_paragraph("當接獲申訴或知悉職場霸凌與性騷擾情事時，企業必須嚴格遵循以下四步驟處理程序：")
    
    table_steps = doc.add_table(rows=1, cols=3)
    table_steps.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_steps.autofit = False
    
    hdr_s = table_steps.rows[0].cells
    hdr_s_titles = ["處置步驟", "核心作業項目", "具體執行內容與法規要求"]
    hdr_s_widths = [Inches(1.5), Inches(1.8), Inches(3.6)]
    
    for i, title in enumerate(hdr_s_titles):
        cell = hdr_s[i]
        cell.width = hdr_s_widths[i]
        set_cell_background(cell, "002060")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9.5)

    steps_data = [
        (
            "步驟一\n即時保護與隔離",
            "立即採取有效之糾正與補救措施",
            "接獲申訴或知悉情事時，雇主應立即啟動保護機制，調整工作排班、座位或進行適當職務隔離，避免申訴人再次遭受冒犯或二次報復。"
        ),
        (
            "步驟二\n獨立調查與申訴處理",
            "啟動獨立調查程序與性平委員會",
            "成立獨立申訴調查小組（外聘專家學者或性平委員比例應符合法定標準），給予雙方充分陳述意見之機會，全程客觀保密處理。"
        ),
        (
            "步驟三\n決議懲處與個案關懷",
            "作成調查報告、懲處與關懷資源",
            "依調查結果進行事實認定與權責懲處（申誡、記過、解僱等），並提供申訴人必要之心理諮商、醫療輔導與法律協助資源，依規通報主管機關。"
        ),
        (
            "步驟四\n宣導教育與滾動追蹤",
            "全員定期宣導與組織文化維護",
            "針對調查發現之組織漏洞進行宣導，定期實施性騷擾與霸凌防治教育訓練，並持續追蹤當事人工作狀況，杜絕冷落或團體排擠。"
        )
    ]

    for r_idx, (s_num, s_title, s_desc) in enumerate(steps_data):
        row_cells = table_steps.add_row().cells
        bg_col = "F9F9F9" if r_idx % 2 == 1 else "FFFFFF"
        
        # Step Num
        cell0 = row_cells[0]
        cell0.width = hdr_s_widths[0]
        set_cell_background(cell0, bg_col)
        set_cell_margins(cell0, top=80, bottom=80, left=80, right=80)
        p0 = cell0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(s_num)
        r0.font.bold = True
        r0.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        
        # Step Title
        cell1 = row_cells[1]
        cell1.width = hdr_s_widths[1]
        set_cell_background(cell1, bg_col)
        set_cell_margins(cell1, top=80, bottom=80, left=80, right=80)
        p1 = cell1.paragraphs[0]
        r1 = p1.add_run(s_title)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
        
        # Step Desc
        cell2 = row_cells[2]
        cell2.width = hdr_s_widths[2]
        set_cell_background(cell2, bg_col)
        set_cell_margins(cell2, top=80, bottom=80, left=80, right=80)
        p2 = cell2.paragraphs[0]
        p2.add_run(s_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    add_card_box("企業法定義務提醒", "當接獲申訴或知悉性騷擾/霸凌情事時，雇主必須立即採取「有效之糾正及補救措施」（包含保護申訴人、啟動獨立調查程序、隔離加害者、給予心理諮商支援），否則依法最高可處新臺幣 100 萬元罰鍰。")

    doc.save(docx_path)
    print(f"Docx generated successfully at: {docx_path}")

def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        word.Quit()
        print(f"PDF generated successfully at: {pdf_path}")
    except Exception as e:
        print(f"Failed to convert PDF: {e}")

if __name__ == "__main__":
    base_dir = r"d:\GOOGLE ANGET"
    docx_file = os.path.join(base_dir, "職場霸凌與性騷擾案例分析與法規辨識手冊.docx")
    pdf_file = os.path.join(base_dir, "職場霸凌與性騷擾案例分析與法規辨識手冊.pdf")
    
    img_factory = r"C:\Users\C606-PC\.gemini\antigravity\brain\44faa0bd-7fa9-4497-a578-37048fb486f2\.user_uploaded\media_1787098103540.jpg"
    img_office = r"C:\Users\C606-PC\.gemini\antigravity\brain\44faa0bd-7fa9-4497-a578-37048fb486f2\.user_uploaded\media_1787098101697.jpg"
    
    generate_harassment_doc(docx_file, img_factory, img_office)
    convert_docx_to_pdf(docx_file, pdf_file)
