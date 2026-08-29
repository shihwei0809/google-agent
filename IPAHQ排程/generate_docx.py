import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Title
title = doc.add_heading('出貨排程管理系統數位化改善提案（含未來 AI 擴展藍圖）', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadata
doc.add_paragraph('提案單位：電子課')
doc.add_paragraph('提案日期：2026/08/29')
doc.add_paragraph('系統名稱：IPA HQ 出貨排程管理系統')
doc.add_paragraph('系統網址：https://tank.shinychem.com.tw')

# Section 1
doc.add_heading('一、現況痛點與耗時分析（業務視角）', level=1)
doc.add_paragraph('電子課每天持續出貨，過去業務人員在處理排程與單據時，面臨大量重複性且高耗時的人工作業：')

doc.add_paragraph('1. 台積電三合一單手工作業（過去每筆需 5 分鐘）：', style='List Bullet')
doc.add_paragraph('每次出貨均需手動開啟 Excel 範本，自行輸入批號、槽車號。', style='List Bullet 2')
doc.add_paragraph('需人工翻查並核對 9 碼的「地點代號」，若輸入錯誤將導致現場無法順利入廠。', style='List Bullet 2')
doc.add_paragraph('無自動化支援前，每張單據的製作、查核與列印，平均耗時約 5 分鐘，極易發生筆誤與重工。', style='List Bullet 2')

doc.add_paragraph('2. CoA 檢驗報告手動比對與登錄（現行每筆需 3 分鐘）：', style='List Bullet')
doc.add_paragraph('目前人員必須一筆一筆慢慢去下載 CoA 圖檔。', style='List Bullet 2')
doc.add_paragraph('接著用肉眼比對、手動複製，再將檢驗數據 K 上系統或單據。', style='List Bullet 2')
doc.add_paragraph('過程繁瑣且容易因為疲勞產生 Key-in 錯誤。', style='List Bullet 2')

# Section 2
doc.add_heading('二、改善內容與系統功能', level=1)
doc.add_paragraph('為解決上述痛點，本次已完成第一階段系統上線，並針對第二階段提出 AI 擴展規劃：')

doc.add_heading('第一階段：台積電三合一單「一鍵自動生成」（已上線）', level=2)
doc.add_paragraph('智慧辨識：系統自動辨識台積電訂單，直接在介面顯示下載按鈕。', style='List Bullet')
doc.add_paragraph('自動容錯查表：後台內建比對邏輯，免除人工翻閱對照表，精準抓出台積電廠區代號。', style='List Bullet')
doc.add_paragraph('一鍵生成與命名：點擊後自動合成批號、槽車號與專屬 QR Code，並嚴格依照規範自動命名（如：2026.8.29. 20P1_S187_台積電槽車barcode三合一單.xlsx），達成「零秒製單」。', style='List Bullet')

doc.add_heading('第二階段藍圖：導入 OCR 技術自動判讀 CoA（建議未來新增）', level=2)
doc.add_paragraph('建議系統未來整合 OCR（光學字元辨識）功能，自動讀取並解析 CoA 圖檔內容。', style='List Bullet')
doc.add_paragraph('自動採檢與驗證：系統自動將擷取到的數據寫入對應欄位，不再需要人工一筆一筆下載、複製、打字，徹底消除這段純勞力密集的工作。', style='List Bullet')


# Section 3
doc.add_heading('三、具體量化效益評估（以每日 50 筆出貨為基準）', level=1)
doc.add_paragraph('由於出貨作業為每日進行（以每月 30 天，全年 365 天計算），透過系統自動化將產生非常可觀的時間複利效益。')

doc.add_heading('1. 單項作業節省時間估算', level=2)
table1 = doc.add_table(rows=1, cols=4)
table1.style = 'Table Grid'
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = '自動化項目'
hdr_cells[1].text = '改善前（純人工）'
hdr_cells[2].text = '改善後（系統/AI）'
hdr_cells[3].text = '每筆省下時間'

row_cells = table1.add_row().cells
row_cells[0].text = '三合一單產製'
row_cells[1].text = '5 分鐘 / 筆'
row_cells[2].text = '3 秒 / 筆'
row_cells[3].text = '約 5 分鐘'

row_cells = table1.add_row().cells
row_cells[0].text = 'CoA OCR 判讀登錄'
row_cells[1].text = '3 分鐘 / 筆'
row_cells[2].text = '自動帶入'
row_cells[3].text = '約 3 分鐘'

row_cells = table1.add_row().cells
row_cells[0].text = '合計'
row_cells[1].text = '8 分鐘 / 筆'
row_cells[2].text = '幾乎為 0'
row_cells[3].text = '8 分鐘 / 筆'

doc.add_heading('2. 整體時間與人力成本節省（量化放大）', level=2)
doc.add_paragraph('若以每天 50 筆常態出貨量計算，自動化帶來的具體節省工時如下：')

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = '時間維度'
hdr_cells[1].text = '三合一單自動化（現已達成）'
hdr_cells[2].text = 'CoA OCR 自動化（未來擴充）'
hdr_cells[3].text = '兩者結合總效益'

row_cells = table2.add_row().cells
row_cells[0].text = '每日節省'
row_cells[1].text = '250 分鐘（約 4.1 小時）'
row_cells[2].text = '150 分鐘（約 2.5 小時）'
row_cells[3].text = '400 分鐘（約 6.6 小時）'

row_cells = table2.add_row().cells
row_cells[0].text = '每月節省'
row_cells[1].text = '7,500 分鐘（約 125 小時）'
row_cells[2].text = '4,500 分鐘（約 75 小時）'
row_cells[3].text = '12,000 分鐘（約 200 小時）'

row_cells = table2.add_row().cells
row_cells[0].text = '每年節省'
row_cells[1].text = '91,250 分鐘（約 1,520 小時）'
row_cells[2].text = '54,750 分鐘（約 912 小時）'
row_cells[3].text = '146,000 分鐘（約 2,433 小時）'

# Section 4
doc.add_heading('四、總結', level=1)
doc.add_paragraph('透過將原本「無三合一單自動化」的 5 分鐘手工作業轉為一鍵生成，單月即可為部門釋放出約 125 小時的工作量。')
doc.add_paragraph('若未來能順利推動「CoA OCR 自動判讀」的開發提案，兩者結合後，等於每個月為部門多出一位全職人員（200小時）的產能。這不僅徹底消除了單據重製與 Key-in 錯誤的隱形成本，更能讓業務人員將這每年省下的 2,400 多個小時，轉而投入於更具價值的核心業務與客戶服務上。')

doc.save('C:/GOOGLE ANGET/IPAHQ排程/提案改善.docx')
print("Docx generated.")
